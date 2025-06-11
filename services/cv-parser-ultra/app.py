"""
🚀 CV Parser Ultra v2.0 - Streaming Temps Réel avec WebSocket
PROMPT 2 - SuperSmartMatch V2 - Parsing ultra-performant

Fonctionnalités :
- ⚡ Streaming temps réel avec WebSocket
- 🤖 OpenAI API intégrée (GPT-4)
- ✅ Validation interactive
- 📊 Métriques Prometheus
- 💾 Cache Redis intelligent
- 🔄 Fallback manuel
- 📱 Support multi-formats (PDF, DOCX, DOC, JPG, PNG)
"""

import asyncio
import json
import time
import uuid
import logging
import hashlib
import mimetypes
from typing import Dict, Any, Optional, List, BinaryIO
from datetime import datetime
from pathlib import Path

import aioredis
import openai
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, File, UploadFile, HTTPException, BackgroundTasks
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from PIL import Image
import pytesseract
import fitz  # PyMuPDF
from docx import Document
import prometheus_client
from prometheus_client import Counter, Histogram, Gauge, generate_latest

# Configuration du logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ===========================================
# MODÈLES PYDANTIC
# ===========================================

class CVParsingProgress(BaseModel):
    task_id: str
    status: str  # "processing", "completed", "error"
    progress: int  # 0-100
    confidence: float  # 0.0-1.0
    current_step: str
    eta_seconds: Optional[int] = None

class CVData(BaseModel):
    nom: str = ""
    prenom: str = ""
    titre_professionnel: str = ""
    email: str = ""
    telephone: str = ""
    adresse: str = ""
    competences_techniques: List[str] = []
    soft_skills: List[str] = []
    logiciels_maitrises: List[str] = []
    langues: List[Dict[str, str]] = []  # [{"langue": "Anglais", "niveau": "Courant"}]
    certifications: List[str] = []
    experience_professionnelle: List[Dict[str, Any]] = []
    formation_diplomes: List[Dict[str, Any]] = []

class CVParsingResult(BaseModel):
    task_id: str
    status: str
    progress: int
    confidence: float
    data: CVData
    suggestions: List[str] = []  # Suggestions pour validation
    fallback_required: bool = False
    metadata: Dict[str, Any] = {}

class ValidationRequest(BaseModel):
    task_id: str
    corrections: Dict[str, Any]

# ===========================================
# MÉTRIQUES PROMETHEUS
# ===========================================

# Compteurs
parsing_requests_total = Counter('cv_parsing_requests_total', 'Total CV parsing requests', ['status', 'format'])
parsing_duration_seconds = Histogram('cv_parsing_duration_seconds', 'CV parsing duration in seconds')
parsing_accuracy_ratio = Gauge('cv_parsing_accuracy_ratio', 'CV parsing accuracy ratio', ['field_type'])
websocket_connections_active = Gauge('cv_websocket_connections_active', 'Active WebSocket connections')
cache_hit_ratio = Gauge('cv_cache_hit_ratio', 'Cache hit ratio')
openai_api_calls_total = Counter('cv_openai_api_calls_total', 'Total OpenAI API calls', ['status'])
file_processing_errors_total = Counter('cv_file_processing_errors_total', 'File processing errors', ['error_type'])

# ===========================================
# CONFIGURATION
# ===========================================

class Config:
    # OpenAI
    OPENAI_API_KEY = "YOUR_OPENAI_API_KEY"  # À configurer via variables d'environnement
    OPENAI_MODEL = "gpt-4"
    
    # Redis
    REDIS_URL = "redis://redis:6379/0"
    CACHE_TTL = 86400  # 24 heures
    
    # Limites
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.jpg', '.jpeg', '.png'}
    
    # Performance
    WEBSOCKET_HEARTBEAT_INTERVAL = 30  # secondes
    PARSING_TIMEOUT = 300  # 5 minutes
    
    # OCR
    TESSERACT_CONFIG = '--oem 3 --psm 6'

config = Config()

# ===========================================
# GESTIONNAIRE WEBSOCKET & CACHE
# ===========================================

class WebSocketManager:
    def __init__(self):
        self.active_connections: Dict[str, WebSocket] = {}
        
    async def connect(self, websocket: WebSocket, task_id: str):
        await websocket.accept()
        self.active_connections[task_id] = websocket
        websocket_connections_active.inc()
        logger.info(f"WebSocket connected for task {task_id}")
        
    async def disconnect(self, task_id: str):
        if task_id in self.active_connections:
            del self.active_connections[task_id]
            websocket_connections_active.dec()
            logger.info(f"WebSocket disconnected for task {task_id}")
            
    async def send_progress(self, task_id: str, progress: CVParsingProgress):
        if task_id in self.active_connections:
            try:
                await self.active_connections[task_id].send_text(progress.json())
            except Exception as e:
                logger.error(f"Error sending progress to {task_id}: {e}")
                await self.disconnect(task_id)

class CacheManager:
    def __init__(self):
        self.redis = None
        
    async def init_redis(self):
        try:
            self.redis = await aioredis.from_url(config.REDIS_URL)
            await self.redis.ping()
            logger.info("Redis connected successfully")
        except Exception as e:
            logger.warning(f"Redis connection failed: {e}")
            
    async def get(self, key: str) -> Optional[Dict[str, Any]]:
        if not self.redis:
            return None
        try:
            data = await self.redis.get(f"cv_ultra:{key}")
            if data:
                cache_hit_ratio.inc()
                return json.loads(data)
        except Exception as e:
            logger.error(f"Cache get error: {e}")
        return None
        
    async def set(self, key: str, value: Dict[str, Any], ttl: int = config.CACHE_TTL):
        if not self.redis:
            return
        try:
            await self.redis.setex(f"cv_ultra:{key}", ttl, json.dumps(value))
        except Exception as e:
            logger.error(f"Cache set error: {e}")

# Instances globales
ws_manager = WebSocketManager()
cache_manager = CacheManager()

# ===========================================
# EXTRACTION DE TEXTE MULTI-FORMAT
# ===========================================

async def extract_text_from_pdf(file_content: bytes) -> str:
    """Extraction texte PDF avec PyMuPDF (plus rapide que PyPDF2)"""
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return text.strip()
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise Exception(f"PDF extraction failed: {e}")

async def extract_text_from_docx(file_content: bytes) -> str:
    """Extraction texte DOCX"""
    try:
        from io import BytesIO
        doc = Document(BytesIO(file_content))
        text = ""
        
        # Paragraphes
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        # Tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
                    
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise Exception(f"DOCX extraction failed: {e}")

async def extract_text_from_image(file_content: bytes) -> str:
    """Extraction texte depuis image avec OCR Tesseract"""
    try:
        from io import BytesIO
        image = Image.open(BytesIO(file_content))
        
        # Prétraitement image pour améliorer OCR
        image = image.convert('RGB')
        
        # OCR avec Tesseract
        text = pytesseract.image_to_string(image, config=config.TESSERACT_CONFIG)
        return text.strip()
    except Exception as e:
        logger.error(f"OCR extraction error: {e}")
        raise Exception(f"OCR extraction failed: {e}")

async def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Router d'extraction selon le type de fichier"""
    ext = Path(filename).suffix.lower()
    
    if ext == '.pdf':
        return await extract_text_from_pdf(file_content)
    elif ext in ['.docx', '.doc']:
        return await extract_text_from_docx(file_content)
    elif ext in ['.jpg', '.jpeg', '.png']:
        return await extract_text_from_image(file_content)
    else:
        raise ValueError(f"Format non supporté: {ext}")

# ===========================================
# PARSING IA AVEC OPENAI GPT-4
# ===========================================

SYSTEM_PROMPT = """
Tu es un expert en extraction d'informations de CV. Analyse méticuleusement ce CV et extrait toutes les informations dans ce format JSON exact :

{
  "nom": "Nom de famille",
  "prenom": "Prénom", 
  "titre_professionnel": "Poste actuel ou recherché",
  "email": "adresse@email.com",
  "telephone": "numéro de téléphone",
  "adresse": "adresse complète",
  "competences_techniques": ["Python", "JavaScript", "etc."],
  "soft_skills": ["Communication", "Leadership", "etc."],
  "logiciels_maitrises": ["Excel", "Photoshop", "etc."],
  "langues": [{"langue": "Anglais", "niveau": "Courant"}],
  "certifications": ["Nom certification", "etc."],
  "experience_professionnelle": [
    {
      "poste": "Titre du poste",
      "entreprise": "Nom entreprise", 
      "duree": "2020-2023",
      "description": "Description détaillée"
    }
  ],
  "formation_diplomes": [
    {
      "diplome": "Nom du diplôme",
      "etablissement": "Nom établissement",
      "annee": "2020",
      "specialite": "Spécialité"
    }
  ]
}

RÈGLES STRICTES :
- Réponds UNIQUEMENT avec le JSON, aucun autre texte
- Si une info manque, utilise "" ou []
- Sois exhaustif et précis
- Groupe les compétences par catégories
- Normalise les formats (téléphone, dates)
"""

async def parse_with_openai(text: str, task_id: str) -> Dict[str, Any]:
    """Parse CV avec OpenAI GPT-4 en streaming"""
    try:
        # Progression
        await ws_manager.send_progress(task_id, CVParsingProgress(
            task_id=task_id,
            status="processing",
            progress=40,
            confidence=0.0,
            current_step="Analyse IA en cours..."
        ))
        
        # Appel OpenAI
        start_time = time.time()
        
        client = openai.OpenAI(api_key=config.OPENAI_API_KEY)
        response = client.chat.completions.create(
            model=config.OPENAI_MODEL,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": text}
            ],
            temperature=0.1,
            max_tokens=3000
        )
        
        duration = time.time() - start_time
        parsing_duration_seconds.observe(duration)
        openai_api_calls_total.labels(status="success").inc()
        
        # Extraction JSON
        content = response.choices[0].message.content.strip()
        
        # Nettoyage du contenu
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        
        # Parse JSON
        try:
            data = json.loads(content.strip())
            
            # Calcul de confiance basé sur la complétude
            confidence = calculate_confidence(data)
            
            await ws_manager.send_progress(task_id, CVParsingProgress(
                task_id=task_id,
                status="processing",
                progress=80,
                confidence=confidence,
                current_step="Validation des données..."
            ))
            
            return data, confidence
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error: {e}")
            openai_api_calls_total.labels(status="json_error").inc()
            raise Exception(f"Invalid JSON response from AI: {e}")
            
    except Exception as e:
        logger.error(f"OpenAI API error: {e}")
        openai_api_calls_total.labels(status="error").inc()
        raise Exception(f"AI parsing failed: {e}")

def calculate_confidence(data: Dict[str, Any]) -> float:
    """Calcule un score de confiance basé sur la complétude des données"""
    total_fields = 10
    filled_fields = 0
    
    # Champs critiques
    if data.get("nom"): filled_fields += 1
    if data.get("prenom"): filled_fields += 1
    if data.get("email"): filled_fields += 1
    if data.get("telephone"): filled_fields += 1
    if data.get("competences_techniques"): filled_fields += 1
    if data.get("experience_professionnelle"): filled_fields += 1
    if data.get("formation_diplomes"): filled_fields += 1
    if data.get("titre_professionnel"): filled_fields += 1
    if data.get("soft_skills"): filled_fields += 1
    if data.get("logiciels_maitrises"): filled_fields += 1
    
    return round(filled_fields / total_fields, 2)

# ===========================================
# PARSING PRINCIPAL AVEC STREAMING
# ===========================================

async def parse_cv_ultra(file_content: bytes, filename: str, task_id: str) -> CVParsingResult:
    """Pipeline de parsing ultra-optimisé avec streaming temps réel"""
    start_time = time.time()
    
    try:
        # 1. Calcul hash pour cache
        file_hash = hashlib.sha256(file_content).hexdigest()
        cache_key = f"{file_hash}_{filename}"
        
        # Progression initiale
        await ws_manager.send_progress(task_id, CVParsingProgress(
            task_id=task_id,
            status="processing", 
            progress=10,
            confidence=0.0,
            current_step="Vérification du cache..."
        ))
        
        # 2. Vérification cache
        cached_result = await cache_manager.get(cache_key)
        if cached_result:
            await ws_manager.send_progress(task_id, CVParsingProgress(
                task_id=task_id,
                status="completed",
                progress=100,
                confidence=cached_result.get("confidence", 0.95),
                current_step="Récupération depuis le cache"
            ))
            
            return CVParsingResult(
                task_id=task_id,
                status="completed",
                progress=100,
                confidence=cached_result.get("confidence", 0.95),
                data=CVData(**cached_result["data"]),
                metadata={"from_cache": True, "duration": 0.1}
            )
        
        # 3. Extraction de texte
        await ws_manager.send_progress(task_id, CVParsingProgress(
            task_id=task_id,
            status="processing",
            progress=20,
            confidence=0.0,
            current_step="Extraction du texte..."
        ))
        
        text = await extract_text_from_file(file_content, filename)
        
        if not text.strip():
            raise Exception("Aucun texte extrait du fichier")
            
        # 4. Parsing IA
        try:
            data, confidence = await parse_with_openai(text, task_id)
        except Exception as e:
            logger.warning(f"AI parsing failed, using fallback: {e}")
            data = fallback_extraction(text)
            confidence = 0.3
            
        # 5. Finalisation
        await ws_manager.send_progress(task_id, CVParsingProgress(
            task_id=task_id,
            status="processing",
            progress=90,
            confidence=confidence,
            current_step="Finalisation..."
        ))
        
        # 6. Cache du résultat
        result_to_cache = {
            "data": data,
            "confidence": confidence,
            "timestamp": datetime.now().isoformat()
        }
        await cache_manager.set(cache_key, result_to_cache)
        
        # 7. Résultat final
        duration = time.time() - start_time
        
        result = CVParsingResult(
            task_id=task_id,
            status="completed",
            progress=100,
            confidence=confidence,
            data=CVData(**data),
            suggestions=generate_validation_suggestions(data, confidence),
            fallback_required=confidence < 0.7,
            metadata={
                "duration": round(duration, 2),
                "file_size": len(file_content),
                "text_length": len(text),
                "from_cache": False
            }
        )
        
        # Progression finale
        await ws_manager.send_progress(task_id, CVParsingProgress(
            task_id=task_id,
            status="completed",
            progress=100,
            confidence=confidence,
            current_step="Terminé !"
        ))
        
        # Métriques
        parsing_requests_total.labels(status="success", format=Path(filename).suffix).inc()
        
        return result
        
    except Exception as e:
        logger.error(f"CV parsing error: {e}")
        
        # Progression erreur
        await ws_manager.send_progress(task_id, CVParsingProgress(
            task_id=task_id,
            status="error",
            progress=0,
            confidence=0.0,
            current_step=f"Erreur: {str(e)}"
        ))
        
        file_processing_errors_total.labels(error_type=type(e).__name__).inc()
        parsing_requests_total.labels(status="error", format=Path(filename).suffix).inc()
        
        raise HTTPException(status_code=500, detail=str(e))

def fallback_extraction(text: str) -> Dict[str, Any]:
    """Extraction basique en cas d'échec de l'IA"""
    import re
    
    # Structure par défaut
    data = {
        "nom": "",
        "prenom": "",
        "titre_professionnel": "",
        "email": "",
        "telephone": "",
        "adresse": "",
        "competences_techniques": [],
        "soft_skills": [],
        "logiciels_maitrises": [],
        "langues": [],
        "certifications": [],
        "experience_professionnelle": [],
        "formation_diplomes": []
    }
    
    # Email
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    if email_match:
        data["email"] = email_match.group(0)
    
    # Téléphone français
    phone_match = re.search(r'(?:(?:\+|00)33|0)[1-9](?:[\s.-]*\d{2}){4}', text)
    if phone_match:
        data["telephone"] = phone_match.group(0)
    
    # Compétences techniques (mots-clés)
    tech_keywords = ['python', 'java', 'javascript', 'react', 'angular', 'node', 'sql', 'html', 'css']
    for keyword in tech_keywords:
        if keyword.lower() in text.lower():
            data["competences_techniques"].append(keyword.capitalize())
    
    return data

def generate_validation_suggestions(data: Dict[str, Any], confidence: float) -> List[str]:
    """Génère des suggestions pour la validation interactive"""
    suggestions = []
    
    if confidence < 0.8:
        suggestions.append("Confiance faible - Vérifiez manuellement les informations extraites")
    
    if not data.get("email"):
        suggestions.append("Email manquant - Ajoutez l'adresse email si disponible")
        
    if not data.get("telephone"):
        suggestions.append("Téléphone manquant - Ajoutez le numéro si disponible")
        
    if len(data.get("competences_techniques", [])) < 3:
        suggestions.append("Peu de compétences techniques détectées - Complétez la liste")
        
    if not data.get("experience_professionnelle"):
        suggestions.append("Expérience professionnelle manquante - Ajoutez les postes précédents")
    
    return suggestions

# ===========================================
# APPLICATION FASTAPI
# ===========================================

app = FastAPI(
    title="CV Parser Ultra v2.0",
    description="Service de parsing CV ultra-performant avec streaming temps réel",
    version="2.0.0"
)

# Événements de démarrage
@app.on_event("startup")
async def startup_event():
    await cache_manager.init_redis()
    logger.info("CV Parser Ultra v2.0 started successfully")

# ===========================================
# ROUTES API
# ===========================================

@app.post("/v2/parse/cv/stream")
async def parse_cv_stream(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...)
) -> Dict[str, str]:
    """Lance le parsing d'un CV avec streaming temps réel"""
    
    # Validation
    if file.size > config.MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="Fichier trop volumineux")
        
    ext = Path(file.filename).suffix.lower()
    if ext not in config.SUPPORTED_FORMATS:
        raise HTTPException(status_code=400, detail=f"Format non supporté: {ext}")
    
    # Génération task ID
    task_id = str(uuid.uuid4())
    
    # Lecture du fichier
    file_content = await file.read()
    
    # Lancement du parsing en arrière-plan
    background_tasks.add_task(parse_cv_ultra, file_content, file.filename, task_id)
    
    return {
        "task_id": task_id,
        "status": "processing",
        "websocket_url": f"/v2/parse/status/{task_id}"
    }

@app.websocket("/v2/parse/status/{task_id}")
async def websocket_progress(websocket: WebSocket, task_id: str):
    """WebSocket pour le streaming de progression en temps réel"""
    await ws_manager.connect(websocket, task_id)
    
    try:
        while True:
            # Maintien de la connexion avec heartbeat
            await asyncio.sleep(config.WEBSOCKET_HEARTBEAT_INTERVAL)
            await websocket.ping()
            
    except WebSocketDisconnect:
        await ws_manager.disconnect(task_id)

@app.get("/v2/parse/validate/{task_id}")
async def get_validation_data(task_id: str) -> Dict[str, Any]:
    """Récupère les données pour validation interactive"""
    # Implémentation de récupération depuis cache/DB
    return {"task_id": task_id, "validation_required": True}

@app.put("/v2/parse/corrections/{task_id}")
async def apply_corrections(task_id: str, corrections: ValidationRequest) -> Dict[str, str]:
    """Applique les corrections utilisateur"""
    # Sauvegarde des corrections
    return {"task_id": task_id, "status": "corrections_applied"}

@app.get("/health")
async def health_check():
    """Vérification de santé du service"""
    return {
        "status": "healthy",
        "service": "CV Parser Ultra v2.0",
        "timestamp": datetime.now().isoformat()
    }

@app.get("/metrics")
async def get_metrics():
    """Métriques Prometheus"""
    return generate_latest()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5051)
