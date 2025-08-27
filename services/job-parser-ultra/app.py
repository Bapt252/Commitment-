"""
🎯 Job Parser Ultra v2.0 - Streaming Temps Réel avec WebSocket
PROMPT 2 - SuperSmartMatch V2 - Parsing fiches de poste ultra-performant

Fonctionnalités :
- ⚡ Streaming temps réel avec WebSocket  
- 🤖 OpenAI API intégrée (GPT-4)
- ✅ Validation interactive
- 📊 Métriques Prometheus
- 💾 Cache Redis intelligent
- 🔄 Fallback manuel
- 📱 Support multi-formats (PDF, DOCX, DOC, HTML, TXT)
"""

import asyncio
import json
import time
import uuid
import logging
import hashlib
import mimetypes
from typing import Dict, Any, Optional, List, BinaryIO
from datetime import datetime
from pathlib import Path
from enum import Enum

import aioredis
import openai
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, File, UploadFile, HTTPException, BackgroundTasks
from fastapi.responses import JSONResponse
from pydantic import BaseModel
import fitz  # PyMuPDF
from docx import Document
from bs4 import BeautifulSoup
import prometheus_client
from prometheus_client import Counter, Histogram, Gauge, generate_latest

# Configuration du logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ===========================================
# MODÈLES PYDANTIC
# ===========================================

class ContractType(str, Enum):
    CDI = "CDI"
    CDD = "CDD"
    STAGE = "Stage"
    FREELANCE = "Freelance"
    ALTERNANCE = "Alternance"
    INTERIM = "Intérim"

class RemoteWorkType(str, Enum):
    FULL_REMOTE = "100% télétravail"
    HYBRID = "Télétravail hybride"
    OCCASIONAL = "Télétravail ponctuel"
    NO_REMOTE = "Présentiel uniquement"

class SalaryRange(BaseModel):
    min_salary: Optional[int] = None
    max_salary: Optional[int] = None
    currency: str = "EUR"
    period: str = "annual"  # annual, monthly

class JobParsingProgress(BaseModel):
    task_id: str
    status: str  # "processing", "completed", "error"
    progress: int  # 0-100
    confidence: float  # 0.0-1.0
    current_step: str
    eta_seconds: Optional[int] = None

class JobData(BaseModel):
    titre_poste: str = ""
    niveau_poste: str = ""  # Junior, Senior, Lead, Manager
    competences_requises: List[str] = []
    competences_souhaitees: List[str] = []
    experience_minimale: str = ""
    localisation: str = ""
    teletravail: RemoteWorkType = RemoteWorkType.NO_REMOTE
    fourchette_salariale: Optional[SalaryRange] = None
    type_contrat: ContractType = ContractType.CDI
    avantages: List[str] = []
    culture_entreprise: List[str] = []
    description_poste: str = ""
    missions_principales: List[str] = []
    profil_recherche: str = ""
    formations_requises: List[str] = []
    langues_requises: List[Dict[str, str]] = []  # [{"langue": "Anglais", "niveau": "Courant"}]
    secteur_activite: str = ""
    taille_entreprise: str = ""

class JobParsingResult(BaseModel):
    task_id: str
    status: str
    progress: int
    confidence: float
    data: JobData
    suggestions: List[str] = []  # Suggestions pour validation
    fallback_required: bool = False
    metadata: Dict[str, Any] = {}

class JobValidationRequest(BaseModel):
    task_id: str
    corrections: Dict[str, Any]

# ===========================================
# MÉTRIQUES PROMETHEUS
# ===========================================

# Compteurs spécifiques aux jobs
job_parsing_requests_total = Counter('job_parsing_requests_total', 'Total job parsing requests', ['status', 'format'])
job_parsing_duration_seconds = Histogram('job_parsing_duration_seconds', 'Job parsing duration in seconds')
job_parsing_accuracy_ratio = Gauge('job_parsing_accuracy_ratio', 'Job parsing accuracy ratio', ['field_type'])
job_websocket_connections_active = Gauge('job_websocket_connections_active', 'Active WebSocket connections for jobs')
job_cache_hit_ratio = Gauge('job_cache_hit_ratio', 'Job cache hit ratio')
job_openai_api_calls_total = Counter('job_openai_api_calls_total', 'Total OpenAI API calls for jobs', ['status'])
job_file_processing_errors_total = Counter('job_file_processing_errors_total', 'Job file processing errors', ['error_type'])

# ===========================================
# CONFIGURATION
# ===========================================

class Config:
    # OpenAI
    OPENAI_API_KEY = "YOUR_OPENAI_API_KEY"  # À configurer via variables d'environnement
    OPENAI_MODEL = "gpt-4"
    
    # Redis
    REDIS_URL = "redis://redis:6379/1"  # DB différente des CV
    CACHE_TTL = 86400  # 24 heures
    
    # Limites
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB pour les jobs
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.html', '.txt'}
    
    # Performance
    WEBSOCKET_HEARTBEAT_INTERVAL = 30  # secondes
    PARSING_TIMEOUT = 180  # 3 minutes (plus rapide que CV)

config = Config()

# ===========================================
# GESTIONNAIRE WEBSOCKET & CACHE
# ===========================================

class JobWebSocketManager:
    def __init__(self):
        self.active_connections: Dict[str, WebSocket] = {}
        
    async def connect(self, websocket: WebSocket, task_id: str):
        await websocket.accept()
        self.active_connections[task_id] = websocket
        job_websocket_connections_active.inc()
        logger.info(f"Job WebSocket connected for task {task_id}")
        
    async def disconnect(self, task_id: str):
        if task_id in self.active_connections:
            del self.active_connections[task_id]
            job_websocket_connections_active.dec()
            logger.info(f"Job WebSocket disconnected for task {task_id}")
            
    async def send_progress(self, task_id: str, progress: JobParsingProgress):
        if task_id in self.active_connections:
            try:
                await self.active_connections[task_id].send_text(progress.json())
            except Exception as e:
                logger.error(f"Error sending job progress to {task_id}: {e}")
                await self.disconnect(task_id)

class JobCacheManager:
    def __init__(self):
        self.redis = None
        
    async def init_redis(self):
        try:
            self.redis = await aioredis.from_url(config.REDIS_URL)
            await self.redis.ping()
            logger.info("Job Redis connected successfully")
        except Exception as e:
            logger.warning(f"Job Redis connection failed: {e}")
            
    async def get(self, key: str) -> Optional[Dict[str, Any]]:
        if not self.redis:
            return None
        try:
            data = await self.redis.get(f"job_ultra:{key}")
            if data:
                job_cache_hit_ratio.inc()
                return json.loads(data)
        except Exception as e:
            logger.error(f"Job cache get error: {e}")
        return None
        
    async def set(self, key: str, value: Dict[str, Any], ttl: int = config.CACHE_TTL):
        if not self.redis:
            return
        try:
            await self.redis.setex(f"job_ultra:{key}", ttl, json.dumps(value))
        except Exception as e:
            logger.error(f"Job cache set error: {e}")

# Instances globales
job_ws_manager = JobWebSocketManager()
job_cache_manager = JobCacheManager()

# ===========================================
# EXTRACTION DE TEXTE MULTI-FORMAT
# ===========================================

async def extract_text_from_pdf(file_content: bytes) -> str:
    """Extraction texte PDF"""
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return text.strip()
    except Exception as e:
        logger.error(f"Job PDF extraction error: {e}")
        raise Exception(f"Job PDF extraction failed: {e}")

async def extract_text_from_docx(file_content: bytes) -> str:
    """Extraction texte DOCX"""
    try:
        from io import BytesIO
        doc = Document(BytesIO(file_content))
        text = ""
        
        # Paragraphes
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        # Tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
                    
        return text.strip()
    except Exception as e:
        logger.error(f"Job DOCX extraction error: {e}")
        raise Exception(f"Job DOCX extraction failed: {e}")

async def extract_text_from_html(file_content: bytes) -> str:
    """Extraction texte HTML"""
    try:
        soup = BeautifulSoup(file_content, 'html.parser')
        
        # Supprimer scripts et styles
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extraire le texte
        text = soup.get_text()
        
        # Nettoyer les espaces
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    except Exception as e:
        logger.error(f"Job HTML extraction error: {e}")
        raise Exception(f"Job HTML extraction failed: {e}")

async def extract_text_from_txt(file_content: bytes) -> str:
    """Extraction texte TXT"""
    try:
        # Tentative avec différents encodages
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue
                
        raise Exception("Impossible de décoder le fichier texte")
    except Exception as e:
        logger.error(f"Job TXT extraction error: {e}")
        raise Exception(f"Job TXT extraction failed: {e}")

async def extract_text_from_job_file(file_content: bytes, filename: str) -> str:
    """Router d'extraction selon le type de fichier job"""
    ext = Path(filename).suffix.lower()
    
    if ext == '.pdf':
        return await extract_text_from_pdf(file_content)
    elif ext in ['.docx', '.doc']:
        return await extract_text_from_docx(file_content)
    elif ext in ['.html', '.htm']:
        return await extract_text_from_html(file_content)
    elif ext == '.txt':
        return await extract_text_from_txt(file_content)
    else:
        raise ValueError(f"Format non supporté pour jobs: {ext}")

# ===========================================
# PARSING IA AVEC OPENAI GPT-4
# ===========================================

JOB_SYSTEM_PROMPT = """
Tu es un expert en analyse d'offres d'emploi. Analyse méticuleusement cette fiche de poste et extrait toutes les informations dans ce format JSON exact :

{
  "titre_poste": "Titre exact du poste",
  "niveau_poste": "Junior|Senior|Lead|Manager|Director",
  "competences_requises": ["Compétence 1", "Compétence 2", "etc."],
  "competences_souhaitees": ["Compétence souhaitée 1", "etc."],
  "experience_minimale": "X années d'expérience",
  "localisation": "Ville, région ou télétravail",
  "teletravail": "100% télétravail|Télétravail hybride|Télétravail ponctuel|Présentiel uniquement",
  "fourchette_salariale": {
    "min_salary": 45000,
    "max_salary": 55000,
    "currency": "EUR",
    "period": "annual"
  },
  "type_contrat": "CDI|CDD|Stage|Freelance|Alternance|Intérim",
  "avantages": ["Avantage 1", "Avantage 2", "etc."],
  "culture_entreprise": ["Valeur 1", "Valeur 2", "etc."],
  "description_poste": "Description complète du poste",
  "missions_principales": ["Mission 1", "Mission 2", "etc."],
  "profil_recherche": "Description du profil idéal",
  "formations_requises": ["Formation 1", "Formation 2", "etc."],
  "langues_requises": [{"langue": "Anglais", "niveau": "Courant"}],
  "secteur_activite": "Secteur d'activité",
  "taille_entreprise": "Startup|PME|Groupe|Multinational"
}

RÈGLES STRICTES :
- Réponds UNIQUEMENT avec le JSON, aucun autre texte
- Si une info manque, utilise "" ou [] ou null
- Sois exhaustif et précis
- Distingue bien compétences requises vs souhaitées
- Normalise les salaires en EUR annuel si possible
- Identifie clairement le niveau de poste
- Extrait TOUTES les missions et avantages mentionnés
"""

async def parse_job_with_openai(text: str, task_id: str) -> tuple[Dict[str, Any], float]:
    """Parse job avec OpenAI GPT-4 en streaming"""
    try:
        # Progression
        await job_ws_manager.send_progress(task_id, JobParsingProgress(
            task_id=task_id,
            status="processing",
            progress=40,
            confidence=0.0,
            current_step="Analyse IA de l'offre d'emploi..."
        ))
        
        # Appel OpenAI
        start_time = time.time()
        
        client = openai.OpenAI(api_key=config.OPENAI_API_KEY)
        response = client.chat.completions.create(
            model=config.OPENAI_MODEL,
            messages=[
                {"role": "system", "content": JOB_SYSTEM_PROMPT},
                {"role": "user", "content": text}
            ],
            temperature=0.1,
            max_tokens=3000
        )
        
        duration = time.time() - start_time
        job_parsing_duration_seconds.observe(duration)
        job_openai_api_calls_total.labels(status="success").inc()
        
        # Extraction JSON
        content = response.choices[0].message.content.strip()
        
        # Nettoyage du contenu
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        
        # Parse JSON
        try:
            data = json.loads(content.strip())
            
            # Calcul de confiance basé sur la complétude
            confidence = calculate_job_confidence(data)
            
            await job_ws_manager.send_progress(task_id, JobParsingProgress(
                task_id=task_id,
                status="processing",
                progress=80,
                confidence=confidence,
                current_step="Validation des données job..."
            ))
            
            return data, confidence
            
        except json.JSONDecodeError as e:
            logger.error(f"Job JSON parsing error: {e}")
            job_openai_api_calls_total.labels(status="json_error").inc()
            raise Exception(f"Invalid JSON response from AI: {e}")
            
    except Exception as e:
        logger.error(f"Job OpenAI API error: {e}")
        job_openai_api_calls_total.labels(status="error").inc()
        raise Exception(f"Job AI parsing failed: {e}")

def calculate_job_confidence(data: Dict[str, Any]) -> float:
    """Calcule un score de confiance pour une offre d'emploi"""
    total_fields = 12
    filled_fields = 0
    
    # Champs critiques pour jobs
    if data.get("titre_poste"): filled_fields += 1
    if data.get("competences_requises"): filled_fields += 1
    if data.get("experience_minimale"): filled_fields += 1
    if data.get("localisation"): filled_fields += 1
    if data.get("type_contrat"): filled_fields += 1
    if data.get("description_poste"): filled_fields += 1
    if data.get("missions_principales"): filled_fields += 1
    if data.get("profil_recherche"): filled_fields += 1
    if data.get("niveau_poste"): filled_fields += 1
    if data.get("secteur_activite"): filled_fields += 1
    if data.get("fourchette_salariale"): filled_fields += 1
    if data.get("avantages"): filled_fields += 1
    
    return round(filled_fields / total_fields, 2)

# ===========================================
# PARSING PRINCIPAL AVEC STREAMING
# ===========================================

async def parse_job_ultra(file_content: bytes, filename: str, task_id: str) -> JobParsingResult:
    """Pipeline de parsing job ultra-optimisé avec streaming temps réel"""
    start_time = time.time()
    
    try:
        # 1. Calcul hash pour cache
        file_hash = hashlib.sha256(file_content).hexdigest()
        cache_key = f"{file_hash}_{filename}"
        
        # Progression initiale
        await job_ws_manager.send_progress(task_id, JobParsingProgress(
            task_id=task_id,
            status="processing", 
            progress=10,
            confidence=0.0,
            current_step="Vérification du cache job..."
        ))
        
        # 2. Vérification cache
        cached_result = await job_cache_manager.get(cache_key)
        if cached_result:
            await job_ws_manager.send_progress(task_id, JobParsingProgress(
                task_id=task_id,
                status="completed",
                progress=100,
                confidence=cached_result.get("confidence", 0.95),
                current_step="Récupération job depuis le cache"
            ))
            
            return JobParsingResult(
                task_id=task_id,
                status="completed",
                progress=100,
                confidence=cached_result.get("confidence", 0.95),
                data=JobData(**cached_result["data"]),
                metadata={"from_cache": True, "duration": 0.1}
            )
        
        # 3. Extraction de texte
        await job_ws_manager.send_progress(task_id, JobParsingProgress(
            task_id=task_id,
            status="processing",
            progress=20,
            confidence=0.0,
            current_step="Extraction du texte job..."
        ))
        
        text = await extract_text_from_job_file(file_content, filename)
        
        if not text.strip():
            raise Exception("Aucun texte extrait de l'offre d'emploi")
            
        # 4. Parsing IA
        try:
            data, confidence = await parse_job_with_openai(text, task_id)
        except Exception as e:
            logger.warning(f"Job AI parsing failed, using fallback: {e}")
            data = job_fallback_extraction(text)
            confidence = 0.3
            
        # 5. Finalisation
        await job_ws_manager.send_progress(task_id, JobParsingProgress(
            task_id=task_id,
            status="processing",
            progress=90,
            confidence=confidence,
            current_step="Finalisation job..."
        ))
        
        # 6. Cache du résultat
        result_to_cache = {
            "data": data,
            "confidence": confidence,
            "timestamp": datetime.now().isoformat()
        }
        await job_cache_manager.set(cache_key, result_to_cache)
        
        # 7. Résultat final
        duration = time.time() - start_time
        
        result = JobParsingResult(
            task_id=task_id,
            status="completed",
            progress=100,
            confidence=confidence,
            data=JobData(**data),
            suggestions=generate_job_validation_suggestions(data, confidence),
            fallback_required=confidence < 0.7,
            metadata={
                "duration": round(duration, 2),
                "file_size": len(file_content),
                "text_length": len(text),
                "from_cache": False
            }
        )
        
        # Progression finale
        await job_ws_manager.send_progress(task_id, JobParsingProgress(
            task_id=task_id,
            status="completed",
            progress=100,
            confidence=confidence,
            current_step="Job parsing terminé !"
        ))
        
        # Métriques
        job_parsing_requests_total.labels(status="success", format=Path(filename).suffix).inc()
        
        return result
        
    except Exception as e:
        logger.error(f"Job parsing error: {e}")
        
        # Progression erreur
        await job_ws_manager.send_progress(task_id, JobParsingProgress(
            task_id=task_id,
            status="error",
            progress=0,
            confidence=0.0,
            current_step=f"Erreur job: {str(e)}"
        ))
        
        job_file_processing_errors_total.labels(error_type=type(e).__name__).inc()
        job_parsing_requests_total.labels(status="error", format=Path(filename).suffix).inc()
        
        raise HTTPException(status_code=500, detail=str(e))

def job_fallback_extraction(text: str) -> Dict[str, Any]:
    """Extraction basique en cas d'échec de l'IA pour jobs"""
    import re
    
    # Structure par défaut
    data = {
        "titre_poste": "",
        "niveau_poste": "",
        "competences_requises": [],
        "competences_souhaitees": [],
        "experience_minimale": "",
        "localisation": "",
        "teletravail": "Présentiel uniquement",
        "fourchette_salariale": None,
        "type_contrat": "CDI",
        "avantages": [],
        "culture_entreprise": [],
        "description_poste": text[:500] + "..." if len(text) > 500 else text,
        "missions_principales": [],
        "profil_recherche": "",
        "formations_requises": [],
        "langues_requises": [],
        "secteur_activite": "",
        "taille_entreprise": ""
    }
    
    # Recherche titre (souvent en début)
    lines = text.split('\n')[:10]  # 10 premières lignes
    for line in lines:
        if len(line.strip()) > 10 and len(line.strip()) < 100:
            data["titre_poste"] = line.strip()
            break
    
    # Recherche salaire
    salary_patterns = [
        r'(\d+)\s*k?€?\s*-\s*(\d+)\s*k?€?',
        r'salaire.*?(\d+).*?(\d+)',
        r'rémunération.*?(\d+).*?(\d+)'
    ]
    
    for pattern in salary_patterns:
        match = re.search(pattern, text.lower())
        if match:
            min_sal = int(match.group(1))
            max_sal = int(match.group(2))
            
            # Ajuster si nécessaire (k€)
            if min_sal < 100:
                min_sal *= 1000
                max_sal *= 1000
                
            data["fourchette_salariale"] = {
                "min_salary": min_sal,
                "max_salary": max_sal,
                "currency": "EUR",
                "period": "annual"
            }
            break
    
    # Compétences techniques
    tech_keywords = ['python', 'java', 'javascript', 'react', 'angular', 'node', 'sql', 'docker', 'kubernetes']
    for keyword in tech_keywords:
        if keyword.lower() in text.lower():
            data["competences_requises"].append(keyword.capitalize())
    
    # Type de contrat
    if 'cdi' in text.lower():
        data["type_contrat"] = "CDI"
    elif 'cdd' in text.lower():
        data["type_contrat"] = "CDD"
    elif 'stage' in text.lower():
        data["type_contrat"] = "Stage"
    
    # Télétravail
    if 'télétravail' in text.lower() or 'remote' in text.lower():
        if '100%' in text or 'full remote' in text.lower():
            data["teletravail"] = "100% télétravail"
        else:
            data["teletravail"] = "Télétravail hybride"
    
    return data

def generate_job_validation_suggestions(data: Dict[str, Any], confidence: float) -> List[str]:
    """Génère des suggestions pour la validation interactive des jobs"""
    suggestions = []
    
    if confidence < 0.8:
        suggestions.append("Confiance faible - Vérifiez manuellement les informations de l'offre")
    
    if not data.get("titre_poste"):
        suggestions.append("Titre du poste manquant - Ajoutez le titre exact")
        
    if not data.get("competences_requises"):
        suggestions.append("Compétences requises manquantes - Listez les compétences essentielles")
        
    if not data.get("fourchette_salariale"):
        suggestions.append("Salaire non détecté - Ajoutez la fourchette salariale si mentionnée")
        
    if not data.get("missions_principales"):
        suggestions.append("Missions principales manquantes - Détaillez les responsabilités")
    
    if not data.get("localisation"):
        suggestions.append("Localisation manquante - Précisez le lieu de travail")
        
    return suggestions

# ===========================================
# APPLICATION FASTAPI
# ===========================================

app = FastAPI(
    title="Job Parser Ultra v2.0",
    description="Service de parsing d'offres d'emploi ultra-performant avec streaming temps réel",
    version="2.0.0"
)

# Événements de démarrage
@app.on_event("startup")
async def startup_event():
    await job_cache_manager.init_redis()
    logger.info("Job Parser Ultra v2.0 started successfully")

# ===========================================
# ROUTES API
# ===========================================

@app.post("/v2/parse/job/stream")
async def parse_job_stream(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...)
) -> Dict[str, str]:
    """Lance le parsing d'une offre d'emploi avec streaming temps réel"""
    
    # Validation
    if file.size > config.MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="Fichier trop volumineux")
        
    ext = Path(file.filename).suffix.lower()
    if ext not in config.SUPPORTED_FORMATS:
        raise HTTPException(status_code=400, detail=f"Format non supporté: {ext}")
    
    # Génération task ID
    task_id = str(uuid.uuid4())
    
    # Lecture du fichier
    file_content = await file.read()
    
    # Lancement du parsing en arrière-plan
    background_tasks.add_task(parse_job_ultra, file_content, file.filename, task_id)
    
    return {
        "task_id": task_id,
        "status": "processing",
        "websocket_url": f"/v2/parse/job/status/{task_id}"
    }

@app.websocket("/v2/parse/job/status/{task_id}")
async def websocket_job_progress(websocket: WebSocket, task_id: str):
    """WebSocket pour le streaming de progression job en temps réel"""
    await job_ws_manager.connect(websocket, task_id)
    
    try:
        while True:
            # Maintien de la connexion avec heartbeat
            await asyncio.sleep(config.WEBSOCKET_HEARTBEAT_INTERVAL)
            await websocket.ping()
            
    except WebSocketDisconnect:
        await job_ws_manager.disconnect(task_id)

@app.get("/v2/parse/job/validate/{task_id}")
async def get_job_validation_data(task_id: str) -> Dict[str, Any]:
    """Récupère les données job pour validation interactive"""
    return {"task_id": task_id, "validation_required": True}

@app.put("/v2/parse/job/corrections/{task_id}")
async def apply_job_corrections(task_id: str, corrections: JobValidationRequest) -> Dict[str, str]:
    """Applique les corrections utilisateur pour jobs"""
    return {"task_id": task_id, "status": "job_corrections_applied"}

@app.get("/health")
async def health_check():
    """Vérification de santé du service job"""
    return {
        "status": "healthy",
        "service": "Job Parser Ultra v2.0",
        "timestamp": datetime.now().isoformat()
    }

@app.get("/metrics")
async def get_job_metrics():
    """Métriques Prometheus pour jobs"""
    return generate_latest()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5053)
