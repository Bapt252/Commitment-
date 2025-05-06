# Job Parser Service - Service de parsing de fiches de poste

import os
import time
import logging
from typing import Dict, Any, Optional, List, BinaryIO
import tempfile
import json
import traceback
import re

from app.core.config import settings
from app.services.resilience import resilient_openai_call
from app.services.mock_parser import get_mock_job_data

# Setup logging
logger = logging.getLogger(__name__)

def parse_job(file_path: str, file_format: Optional[str] = None) -> Dict[str, Any]:
    """Parse une fiche de poste pour en extraire les informations structurées
    
    Args:
        file_path: Chemin vers le fichier de fiche de poste
        file_format: Format du fichier (.pdf, .docx, etc.)
        
    Returns:
        Dict[str, Any]: Informations structurées extraites de la fiche de poste
    """
    # 1. Déterminer le format si non fourni
    if not file_format and file_path:
        file_format = os.path.splitext(file_path)[1].lower()
    
    # Logging du fichier traité pour le debugging
    logger.info(f"Traitement du fichier: {os.path.basename(file_path)} (format: {file_format})")
    
    try:
        # 2. Extraire le texte de la fiche de poste selon le format
        job_text = extract_text_from_file(file_path, file_format)
        
        # Log de la taille du texte extrait pour debug
        logger.info(f"Texte extrait: {len(job_text)} caractères")
        if len(job_text) < 100:  # Si le texte extrait est très court, c'est probablement un problème
            logger.warning(f"Texte extrait très court ({len(job_text)} caractères), possible problème d'extraction")
            logger.debug(f"Contenu extrait: {job_text}")
        
        # Pré-traitement du texte pour améliorer la détection
        job_text = preprocess_job_text(job_text)
        
        # 3. Utiliser OpenAI pour analyser la fiche de poste
        start_time = time.time()
        
        try:
            # Si USE_MOCK_PARSER est activé, utiliser le mock au lieu de l'API
            if settings.USE_MOCK_PARSER:
                logger.info(f"Utilisation du mock parser (mode de simulation) pour {file_path}")
                parsed_data = get_mock_job_data(job_text, os.path.basename(file_path))
            else:
                # Sinon, utiliser l'API OpenAI
                parsed_data = analyze_job_with_gpt(job_text)
                
                # Post-traitement pour corriger et enrichir les données
                parsed_data = postprocess_job_data(parsed_data, job_text)
        except Exception as e:
            logger.error(f"Erreur lors de l'analyse de la fiche de poste: {str(e)}. Fallback sur le mock parser.")
            logger.error(f"Stacktrace: {traceback.format_exc()}")
            # En cas d'erreur, utiliser le mock parser comme fallback
            parsed_data = get_mock_job_data(job_text, os.path.basename(file_path))
        
        processing_time = time.time() - start_time
        logger.info(f"Fiche de poste parsée en {processing_time:.2f} secondes")
        
        # 4. Ajouter des métadonnées au résultat
        result = {
            "processing_time": processing_time,
            "parsed_at": time.time(),
            "file_format": file_format,
            "model": "mock" if settings.USE_MOCK_PARSER else settings.OPENAI_MODEL,
            "data": parsed_data
        }
        
        return result
    except Exception as e:
        logger.error(f"Erreur pendant le parsing de la fiche de poste {os.path.basename(file_path)}: {str(e)}")
        logger.error(f"Stacktrace: {traceback.format_exc()}")
        
        # Retourner un résultat avec l'erreur mais une structure minimale valide
        return {
            "processing_time": 0,
            "parsed_at": time.time(),
            "file_format": file_format,
            "model": "error",
            "error": str(e),
            "data": {
                "title": "",
                "company": "",
                "location": "",
                "contract_type": "",
                "required_skills": [],
                "preferred_skills": [],
                "responsibilities": [],
                "requirements": [],
                "benefits": []
            }
        }

def preprocess_job_text(text: str) -> str:
    """Prétraite le texte de la fiche de poste pour améliorer la qualité de l'analyse"""
    if not text:
        return ""
        
    # Remplacer les séquences de plusieurs espaces par un seul
    text = re.sub(r' +', ' ', text)
    
    # Remplacer les séquences de plusieurs lignes vides par une seule
    text = re.sub(r'\n+', '\n', text)
    
    return text

def postprocess_job_data(data: Dict[str, Any], original_text: str) -> Dict[str, Any]:
    """Post-traitement des données extraites pour corriger et enrichir les informations"""
    if not data:
        return {}
    
    # S'assurer que tous les champs requis existent
    required_fields = [
        "title", "company", "location", "contract_type", 
        "required_skills", "preferred_skills", "responsibilities", 
        "requirements", "benefits"
    ]
    
    for field in required_fields:
        if field not in data:
            if field in ["required_skills", "preferred_skills", "responsibilities", "requirements", "benefits"]:
                data[field] = []
            else:
                data[field] = ""
    
    # Extraire le type de contrat s'il n'est pas détecté
    if not data.get("contract_type"):
        contract_type = extract_contract_type(original_text)
        if contract_type:
            data["contract_type"] = contract_type
    
    # Extraire les compétences requises si aucune n'a été détectée
    if not data.get("required_skills"):
        skills = extract_skills_from_text(original_text)
        if skills:
            data["required_skills"] = skills
    
    return data

def extract_contract_type(text: str) -> str:
    """Extrait le type de contrat du texte de la fiche de poste"""
    contract_patterns = [
        r'(?i)\b(CDI|contrat à durée indéterminée)\b',
        r'(?i)\b(CDD|contrat à durée déterminée)\b',
        r'(?i)\b(stage|internship)\b',
        r'(?i)\b(freelance|indépendant)\b',
        r'(?i)\b(alternance|apprentissage)\b',
        r'(?i)\b(temps partiel|part[ -]time)\b',
        r'(?i)\b(temps plein|full[ -]time)\b'
    ]
    
    for pattern in contract_patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0).capitalize()
    
    return ""

def extract_skills_from_text(text: str) -> List[str]:
    """Extrait des compétences potentielles du texte de la fiche de poste"""
    skills = []
    
    # Liste de compétences courantes en finance/comptabilité
    common_skills = [
        'Comptabilité générale', 'Comptabilité analytique', 'Comptabilité clients', 
        'Comptabilité fournisseurs', 'Fiscalité', 'Audit', 'Contrôle de gestion',
        'Gestion de trésorerie', 'Finance d\'entreprise', 'Normes IFRS', 'Normes US GAAP',
        'Consolidation', 'Reporting', 'Budget', 'Prévisions', 'Analyse financière',
        'Clôture comptable', 'Rapprochement bancaire', 'Liasse fiscale', 'Bilan',
        'SAP', 'Oracle', 'Sage', 'Excel', 'Power BI', 'Anglais'
    ]
    
    # Rechercher les compétences dans le texte
    for skill in common_skills:
        if re.search(r'\b' + re.escape(skill) + r'\b', text, re.IGNORECASE):
            skills.append(skill)
    
    return skills

def extract_text_from_file(file_path: str, file_format: Optional[str] = None) -> str:
    """Extrait le texte d'un fichier
    
    Args:
        file_path: Chemin vers le fichier
        file_format: Format du fichier
        
    Returns:
        str: Texte extrait du fichier
    """
    logger.info(f"Extraction du texte depuis {file_path} (format: {file_format})")
    
    # Déterminer le format si non fourni
    if not file_format:
        file_format = os.path.splitext(file_path)[1].lower()
    
    # Vérifier si le fichier existe pour éviter des erreurs
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Le fichier n'existe pas: {file_path}")
    
    # Log de la taille du fichier
    file_size = os.path.getsize(file_path)
    logger.info(f"Taille du fichier: {file_size / 1024:.2f} KB")
    
    # Extraction selon le format avec gestion d'erreur améliorée
    try:
        if file_format.lower() in [".pdf", ".PDF"]:
            return extract_text_from_pdf(file_path)
        elif file_format.lower() in [".docx", ".DOCX"]:
            return extract_text_from_docx(file_path)
        elif file_format.lower() in [".doc", ".DOC"]:
            return extract_text_from_doc(file_path)
        elif file_format.lower() in [".txt", ".TXT", ".text"]:
            return extract_text_from_txt(file_path)
        elif file_format.lower() in [".rtf", ".RTF"]:
            return extract_text_from_rtf(file_path)
        else:
            # Tenter une extraction générique pour les formats non reconnus
            logger.warning(f"Format non reconnu: {file_format}. Tentative d'extraction générique.")
            return extract_text_generic(file_path)
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte ({file_format}): {str(e)}")
        logger.error(f"Stacktrace: {traceback.format_exc()}")
        # En cas d'échec, tenter une extraction générique alternative
        try:
            logger.info("Tentative d'extraction alternative...")
            return extract_text_generic(file_path)
        except Exception as alt_e:
            logger.error(f"Échec de l'extraction alternative: {str(alt_e)}")
            # Retourner une chaîne vide mais non None pour éviter des erreurs en aval
            return f"Échec d'extraction du texte. Format: {file_format}. Erreur: {str(e)}"

def extract_text_generic(file_path: str) -> str:
    """Méthode d'extraction générique qui tente plusieurs approches"""
    logger.info(f"Extraction générique pour {file_path}")
    
    # Essayer des méthodes alternatives d'extraction
    try:
        # Tenter avec textract qui supporte de nombreux formats
        try:
            import textract
            text = textract.process(file_path).decode('utf-8')
            if text and len(text) > 100:  # Vérifier qu'on a extrait quelque chose de significatif
                return text
        except:
            logger.warning("Échec de l'extraction avec textract")
        
        # Tenter avec pdfplumber (autre bibliothèque pour PDF)
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
                if text:
                    return text
        except:
            logger.warning("Échec de l'extraction avec pdfplumber")
        
        # Utiliser notre propre méthode PDF comme dernier recours
        try:
            return extract_text_from_pdf(file_path)
        except:
            logger.warning("Échec de l'extraction PDF standard")
        
        # Si tout échoue, tenter une lecture binaire simple
        with open(file_path, 'rb') as f:
            content = f.read()
            try:
                # Tenter plusieurs encodages
                for encoding in ['utf-8', 'latin-1', 'windows-1252', 'ascii']:
                    try:
                        return content.decode(encoding)
                    except:
                        continue
            except:
                pass
        
        return "Extraction de texte échouée pour ce document."
    except Exception as e:
        logger.error(f"Toutes les méthodes d'extraction ont échoué: {str(e)}")
        return "Échec de toutes les méthodes d'extraction de texte."

def extract_text_from_pdf(file_path: str) -> str:
    """Extrait le texte d'un fichier PDF avec meilleure gestion d'erreurs"""
    try:
        logger.info(f"Tentative d'extraction PDF depuis {file_path}")
        
        # Première tentative avec PyPDF2
        try:
            from PyPDF2 import PdfReader
            
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                text = ""
                for page in reader.pages:
                    extracted_text = page.extract_text()
                    if extracted_text:
                        text += extracted_text + "\n"
                
                if text.strip():  # Vérifier que le texte extrait n'est pas vide
                    logger.info(f"Extraction PyPDF2 réussie: {len(text)} caractères")
                    return text
                else:
                    logger.warning("PyPDF2 n'a pas extrait de texte - le PDF pourrait être scanné ou contenir des images")
        except Exception as e:
            logger.warning(f"Échec de l'extraction avec PyPDF2: {str(e)}")
        
        # Deuxième tentative avec pdfminer.six
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            text = pdfminer_extract(file_path)
            if text.strip():
                logger.info(f"Extraction pdfminer.six réussie: {len(text)} caractères")
                return text
            else:
                logger.warning("pdfminer.six n'a pas extrait de texte")
        except Exception as e:
            logger.warning(f"Échec de l'extraction avec pdfminer.six: {str(e)}")
        
        # Troisième tentative avec OCR si les autres méthodes échouent
        try:
            import pytesseract
            from PIL import Image
            import pdf2image
            
            logger.info("Tentative d'extraction via OCR (conversion PDF en images puis OCR)")
            pages = pdf2image.convert_from_path(file_path)
            text = ""
            for page in pages:
                text += pytesseract.image_to_string(page) + "\n"
            
            if text.strip():
                logger.info(f"Extraction OCR réussie: {len(text)} caractères")
                return text
        except Exception as e:
            logger.warning(f"Échec de l'extraction avec OCR: {str(e)}")
        
        # Si toutes les méthodes échouent
        return "Texte non extractible de ce PDF. Il pourrait s'agir d'un document scanné ou protégé."
        
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte PDF: {str(e)}")
        raise

def extract_text_from_docx(file_path: str) -> str:
    """Extrait le texte d'un fichier DOCX"""
    try:
        import docx
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extraire également les tableaux qui peuvent contenir des informations importantes
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return text
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte DOCX: {str(e)}")
        raise

def extract_text_from_doc(file_path: str) -> str:
    """Extrait le texte d'un fichier DOC (ancien format Word)"""
    try:
        # Plusieurs approches possibles
        methods_tried = []
        
        # Essayer avec textract (nécessite installation système)
        try:
            import textract
            text = textract.process(file_path).decode('utf-8')
            methods_tried.append("textract")
            if text:
                return text
        except ImportError:
            logger.warning("Module textract non disponible pour l'extraction DOC")
        except Exception as e:
            logger.warning(f"Échec de l'extraction DOC avec textract: {str(e)}")
        
        # Essayer avec antiword
        try:
            import subprocess
            result = subprocess.run(['antiword', file_path], stdout=subprocess.PIPE)
            text = result.stdout.decode('utf-8')
            methods_tried.append("antiword")
            if text:
                return text
        except Exception as e:
            logger.warning(f"Échec de l'extraction DOC avec antiword: {str(e)}")
        
        # Essayer avec pywin32 (Windows uniquement)
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)
            text = doc.Content.Text
            doc.Close()
            word.Quit()
            methods_tried.append("win32com")
            if text:
                return text
        except Exception as e:
            logger.warning(f"Échec de l'extraction DOC avec win32com: {str(e)}")

        if not methods_tried:
            raise NotImplementedError("Aucune méthode d'extraction DOC n'a fonctionné")
        else:
            raise Exception(f"Échec de l'extraction DOC avec les méthodes: {', '.join(methods_tried)}")
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte DOC: {str(e)}")
        raise

def extract_text_from_txt(file_path: str) -> str:
    """Extrait le texte d'un fichier texte"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Essayer avec des encodages alternatifs
        for encoding in ['latin-1', 'windows-1252', 'iso-8859-1', 'cp1252']:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        # Si tous les encodages échouent, essayer en mode binaire
        with open(file_path, 'rb') as file:
            content = file.read()
            return str(content)
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte TXT: {str(e)}")
        raise

def extract_text_from_rtf(file_path: str) -> str:
    """Extrait le texte d'un fichier RTF"""
    try:
        # Essayer avec striprtf
        try:
            from striprtf.striprtf import rtf_to_text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            return rtf_to_text(content)
        except ImportError:
            # Fallback à textract
            import textract
            text = textract.process(file_path).decode('utf-8')
            return text
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte RTF: {str(e)}")
        raise

def analyze_job_with_gpt(job_text: str) -> Dict[str, Any]:
    """Analyse une fiche de poste avec GPT pour en extraire les informations structurées
    
    Args:
        job_text: Texte de la fiche de poste
        
    Returns:
        Dict[str, Any]: Informations structurées extraites de la fiche de poste
    """
    logger.info(f"Analyse de la fiche de poste avec {settings.OPENAI_MODEL} (longueur: {len(job_text)} caractères)")
    
    # Si le texte est vide ou trop court, retourner une structure vide
    if not job_text or len(job_text) < 50:
        logger.warning(f"Texte de fiche de poste trop court ({len(job_text)} caractères), impossible d'analyser")
        return {
            "title": "",
            "company": "",
            "location": "",
            "contract_type": "",
            "required_skills": [],
            "preferred_skills": [],
            "responsibilities": [],
            "requirements": [],
            "benefits": []
        }
    
    # Si le texte est trop long, le tronquer pour éviter des problèmes avec l'API
    max_tokens = 15000  # Approximativement 15000 caractères
    if len(job_text) > max_tokens:
        logger.warning(f"Fiche de poste trop longue ({len(job_text)} caractères), troncature à {max_tokens} caractères")
        job_text = job_text[:max_tokens] + "...[texte tronqué]"
    
    # Définir le prompt pour l'extraction d'information structurée
    prompt = f"""
Tu es un expert en analyse de fiches de poste pour l'industrie du recrutement.
Tu dois extraire avec précision toutes les informations importantes d'une fiche de poste.

INSTRUCTIONS IMPÉRATIVES:
1. Extrait UNIQUEMENT les informations réellement présentes dans la fiche de poste.
2. Ne génère JAMAIS d'informations fictives.
3. Pour tout champ non présent dans la fiche de poste, renvoie une valeur vide (chaîne vide ou tableau vide).
4. Sois particulièrement attentif au titre du poste, à l'entreprise, au lieu de travail et au type de contrat.
5. Différencie bien les compétences requises des compétences souhaitées.

EXTRACTION DEMANDÉE:
Retourne un JSON avec cette structure précise et TOUS ces champs, même vides:

{{
  "title": "",           // Titre du poste
  "company": "",        // Nom de l'entreprise
  "location": "",       // Lieu de travail
  "contract_type": "",  // Type de contrat (CDI, CDD, freelance, etc.)
  "required_skills": [  // Compétences requises (obligatoires)
    "Compétence 1",
    "Compétence 2"
  ],
  "preferred_skills": [ // Compétences souhaitées (optionnelles)
    "Compétence A",
    "Compétence B"
  ],
  "responsibilities": [  // Missions et responsabilités
    "Responsabilité 1",
    "Responsabilité 2"
  ],
  "requirements": [     // Prérequis (formation, expérience, etc.)
    "Prérequis 1",
    "Prérequis 2"
  ],
  "benefits": [        // Avantages proposés
    "Avantage 1",
    "Avantage 2"
  ],
  "salary_range": "",  // Fourchette de salaire (si mentionnée)
  "remote_policy": "", // Politique de télétravail (si mentionnée)
  "application_process": "", // Processus de candidature
  "company_description": ""  // Description de l'entreprise
}}

FICHE DE POSTE À ANALYSER:
{job_text}

IMPORTANT: Tu dois retourner UNIQUEMENT le JSON avec toutes les informations récupérées de la fiche de poste, sans aucun texte d'introduction ou commentaire.
"""
    
    try:
        # Appel résilient à OpenAI (avec circuit breaker et retry)
        response_text = resilient_openai_call(
            prompt=prompt, 
            model=settings.OPENAI_MODEL,
            temperature=0.1,
            max_tokens=4000
        )
        
        logger.debug(f"Réponse brute d'OpenAI: {response_text}")
        
        # Parser la réponse JSON
        try:
            # Nettoyage du texte pour s'assurer qu'il ne contient que du JSON
            import re
            json_pattern = r'(\{[\s\S]*\})' 
            match = re.search(json_pattern, response_text)
            if match:
                json_str = match.group(1)
                try:
                    parsed_result = json.loads(json_str)
                    logger.info("Parsing JSON réussi")
                    return parsed_result
                except json.JSONDecodeError as json_err:
                    logger.error(f"Erreur de décodage JSON après extraction: {str(json_err)}")
            
            # Essayer de parser directement si l'extraction a échoué
            try:
                parsed_result = json.loads(response_text)
                return parsed_result
            except json.JSONDecodeError:
                pass
                
            # Tentative de nettoyage et d'extraction plus agressive
            clean_text = response_text.strip()
            if clean_text.startswith("```json"):
                clean_text = clean_text[7:]
            if clean_text.endswith("```"):
                clean_text = clean_text[:-3]
            clean_text = clean_text.strip()
            
            try:
                return json.loads(clean_text)
            except json.JSONDecodeError:
                logger.error("Échec de la tentative de nettoyage JSON")
                
                # En dernier recours, renvoyer un dictionnaire par défaut
                return {
                    "error": "Format JSON invalide dans la réponse",
                    "title": "",
                    "company": "",
                    "location": "",
                    "contract_type": "",
                    "required_skills": [],
                    "preferred_skills": [],
                    "responsibilities": [],
                    "requirements": [],
                    "benefits": []
                }
        except Exception as parse_err:
            logger.error(f"Erreur lors du parsing de la réponse: {str(parse_err)}")
            # Structure de retour par défaut en cas d'erreur
            return {
                "error": f"Erreur de parsing: {str(parse_err)}",
                "title": "",
                "company": "",
                "location": "",
                "contract_type": "",
                "required_skills": [],
                "preferred_skills": [],
                "responsibilities": [],
                "requirements": [],
                "benefits": []
            }
    
    except Exception as e:
        logger.error(f"Erreur lors de l'analyse de la fiche de poste avec GPT: {str(e)}")
        # Structure de retour par défaut en cas d'erreur
        return {
            "error": f"Erreur d'analyse: {str(e)}",
            "title": "",
            "company": "",
            "location": "",
            "contract_type": "",
            "required_skills": [],
            "preferred_skills": [],
            "responsibilities": [],
            "requirements": [],
            "benefits": []
        }
