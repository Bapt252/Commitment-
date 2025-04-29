import os
import json
import hashlib
import logging
import time
from typing import Dict, Any, BinaryIO, Optional, Tuple, List, Union

import redis
import openai
from PyPDF2 import PdfReader
from docx import Document
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

from app.models.cv_model import CVModel

# Configuration du logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Configuration de l'API OpenAI - Support pour les deux noms de variables
openai.api_key = os.getenv("OPENAI_API_KEY") or os.getenv("OPENAI")
if not openai.api_key:
    logger.warning("OpenAI API key not found in environment variables. API calls will fail.")

# Configuration Redis
REDIS_HOST = os.getenv("REDIS_HOST", "redis")
REDIS_PORT = int(os.getenv("REDIS_PORT", 6379))
REDIS_DB = int(os.getenv("REDIS_DB", 0))
REDIS_EXPIRY = int(os.getenv("REDIS_EXPIRY", 86400 * 30))  # 30 jours par défaut

# Connexion Redis avec gestion d'erreur
try:
    redis_client = redis.Redis(
        host=REDIS_HOST, 
        port=REDIS_PORT, 
        db=REDIS_DB,
        decode_responses=False  # Conserver les données binaires pour le cache
    )
    redis_client.ping()  # Vérifie si la connexion est établie
    logger.info("Connected to Redis successfully")
except redis.ConnectionError as e:
    logger.warning(f"Redis connection failed: {e}. Cache will be disabled.")
    redis_client = None

# Constantes
CACHE_PREFIX = "cv_parse:"
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc'}

# Prompt système optimisé avec focus sur les expériences professionnelles
SYSTEM_PROMPT = """
Tu es un assistant spécialisé dans l'extraction précise d'informations de CV pour le projet Commitment. Ta tâche principale est d'analyser le CV fourni et d'en extraire les informations structurées selon un format JSON spécifique.

## RÈGLES GÉNÉRALES
- Extrais UNIQUEMENT les informations explicitement présentes dans le CV.
- N'invente JAMAIS d'informations.
- Utilise une chaîne vide ou liste vide pour tout champ sans information explicite.
- Respecte scrupuleusement le format JSON demandé.

## FORMAT DE SORTIE JSON
{
  "nom": string,
  "prenom": string,
  "poste": string,
  "adresse": string,
  "email": string,
  "telephone": string,
  "competences": [string],
  "logiciels": [string],
  "soft_skills": [string],
  "experiences": [
    {
      "titre": string,
      "entreprise": string,
      "lieu": string,
      "date_debut": string,  // Format ISO "YYYY-MM"
      "date_fin": string | "Present",  // Format ISO "YYYY-MM" ou "Present"
      "description": string,
      "responsabilites": [string],  // Liste des responsabilités spécifiques
      "realisations": [string],  // Liste des réalisations quantifiables
      "technologies": [string],  // Technologies mentionnées dans cette expérience
      "taille_equipe": number | "",
      "type_contrat": string | ""  // CDI, CDD, Freelance, Stage, etc.
    }
  ],
  "formation": [
    {
      "diplome": string,
      "etablissement": string,
      "lieu": string,
      "date_debut": string,  // Format ISO "YYYY-MM"
      "date_fin": string,    // Format ISO "YYYY-MM"
      "description": string,
      "distinctions": [string]
    }
  ],
  "certifications": [
    {
      "nom": string,
      "organisme": string,
      "date_obtention": string,
      "date_expiration": string | "No Expiration"
    }
  ]
}

## INSTRUCTIONS SPÉCIALES POUR LES EXPÉRIENCES PROFESSIONNELLES
1. PRÉCISION TEMPORELLE:
   - Convertis toutes les dates au format "YYYY-MM"
   - Ex: "Janvier 2022" → "2022-01", "Mars 2020" → "2020-03"
   - Si seule l'année est mentionnée, utilise le mois de janvier: "2019" → "2019-01"
   - Pour les expériences en cours, utilise exactement "Present" pour le champ date_fin

2. DÉCOMPOSITION STRUCTURÉE:
   - Distingue clairement entre:
     * "description": résumé général du poste ou de la mission
     * "responsabilites": tâches et activités régulières (utilise des verbes d'action)
     * "realisations": résultats concrets et mesurables (avec des métriques si possible)

3. EXTRACTION TECHNOLOGIQUE:
   - Identifie toutes les technologies, frameworks, outils, méthodologies mentionnés dans chaque expérience
   - Standardise les noms (ex: "React.js" → "React", "Tensorflow" → "TensorFlow")

4. ANALYSE CONTEXTUELLE:
   - Détecte le type de contrat (CDI, CDD, Freelance, Stage) à partir du contexte
   - Identifie la taille d'équipe lorsqu'elle est mentionnée, même indirectement
   - Extrais le lieu précis de l'expérience (ville et pays si disponible)

5. ORDRE CHRONOLOGIQUE:
   - Place les expériences dans l'ordre chronologique inversé (la plus récente en premier)

Réponds UNIQUEMENT avec un objet JSON valide correspondant au format demandé, sans aucun texte avant ou après.
"""

# Prompt de raffinement pour les expériences professionnelles
EXPERIENCE_REFINEMENT_PROMPT = """
Tu es un expert en analyse d'expériences professionnelles issues de CV.

Analyse et améliore les expériences extraites selon ces critères:
1. Sépare clairement les responsabilités des réalisations concrètes
2. Extrait toutes les technologies mentionnées dans chaque expérience
3. Identifie les métriques et résultats quantifiables pour les réalisations
4. Normalise les dates au format YYYY-MM
5. Identifie la taille d'équipe et le type de contrat si indiqués

Réponds UNIQUEMENT avec un tableau JSON des expériences améliorées selon le format initial.
"""

class CVParserError(Exception):
    """Exception personnalisée pour les erreurs de parsing de CV"""
    pass

class FileExtractionError(CVParserError):
    """Exception levée quand l'extraction du texte d'un fichier échoue"""
    pass

class OpenAIError(CVParserError):
    """Exception levée quand l'appel à l'API OpenAI échoue"""
    pass

class JSONParsingError(CVParserError):
    """Exception levée quand le parsing du JSON échoue"""
    pass

def calculate_file_hash(file: BinaryIO) -> str:
    """Calcule un hash SHA-256 du contenu du fichier pour l'utiliser comme clé de cache."""
    file.seek(0)
    file_content = file.read()
    file_hash = hashlib.sha256(file_content).hexdigest()
    file.seek(0)  # Réinitialiser le curseur pour une utilisation ultérieure
    return file_hash

def extract_text_from_pdf(file: BinaryIO) -> str:
    """Extrait le texte d'un fichier PDF."""
    try:
        pdf_reader = PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:  # Vérifier que le texte n'est pas None
                text += page_text + "\n"
        
        if not text.strip():
            logger.warning(f"Extracted empty text from PDF file")
            
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise FileExtractionError(f"PDF extraction error: {str(e)}")

def extract_text_from_docx(file: BinaryIO) -> str:
    """Extrait le texte d'un fichier DOCX."""
    try:
        doc = Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text:  # Vérifier que le texte n'est pas None
                text += paragraph.text + "\n"
        
        # Extraire également le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + "\n"
        
        if not text.strip():
            logger.warning(f"Extracted empty text from DOCX file")
            
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise FileExtractionError(f"DOCX extraction error: {str(e)}")

def extract_text(file: BinaryIO, file_extension: str) -> str:
    """Extrait le texte d'un fichier selon son extension."""
    # Reset file cursor position to beginning
    file.seek(0)
    
    if file_extension.lower() == '.pdf':
        return extract_text_from_pdf(file)
    elif file_extension.lower() in ['.docx', '.doc']:
        return extract_text_from_docx(file)
    else:
        raise ValueError(f"Format de fichier non supporté: {file_extension}")

def get_cached_result(cache_key: str) -> Optional[Dict[str, Any]]:
    """
    Récupère un résultat mis en cache dans Redis
    
    Args:
        cache_key: Clé de cache
        
    Returns:
        Le résultat en cache si présent, None sinon
    """
    if redis_client is None:
        return None
        
    try:
        cached_data = redis_client.get(cache_key)
        if cached_data:
            return json.loads(cached_data)
        return None
    except Exception as e:
        logger.warning(f"Error retrieving from cache: {str(e)}")
        return None

def cache_result(cache_key: str, result: Dict[str, Any]) -> bool:
    """
    Met en cache un résultat dans Redis
    
    Args:
        cache_key: Clé de cache
        result: Résultat à mettre en cache
        
    Returns:
        True si le caching a réussi, False sinon
    """
    if redis_client is None:
        return False
        
    try:
        redis_client.set(
            cache_key,
            json.dumps(result),
            ex=REDIS_EXPIRY
        )
        return True
    except Exception as e:
        logger.warning(f"Error caching result: {str(e)}")
        return False

@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=2, max=10),
    retry=retry_if_exception_type(OpenAIError),
    reraise=True
)
def parse_cv_with_openai(cv_text: str) -> Dict[str, Any]:
    """Parse un CV en utilisant l'API OpenAI GPT-4o-mini avec mécanisme de retry."""
    # Vérifier la clé API
    if not openai.api_key:
        raise OpenAIError("La clé API OpenAI n'est pas configurée dans les variables d'environnement.")
    
    # Limiter la taille du texte pour l'API
    max_tokens = 16000  # Limite prudente pour gpt-4o-mini
    if len(cv_text) > max_tokens:
        logger.warning(f"CV text too long ({len(cv_text)} chars), truncating to {max_tokens} chars")
        cv_text = cv_text[:max_tokens]
    
    # Appeler l'API OpenAI
    try:
        start_time = time.time()
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": cv_text}
            ],
            temperature=0.1,  # Température basse pour plus de cohérence
            max_tokens=3000,  # Augmenté pour traiter plus de détails
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )
        
        execution_time = time.time() - start_time
        logger.info(f"OpenAI API request completed in {execution_time:.2f} seconds")
        
        # Extraire le contenu de la réponse
        content = response.choices[0].message.content
        
        # Nettoyer le contenu au cas où GPT aurait ajouté des backticks ou commentaires
        content = content.strip()
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()
        
        try:
            # Essayer d'extraire uniquement la partie JSON si le modèle a inclus d'autres textes
            json_start = content.find('{')
            json_end = content.rfind('}')
            
            if json_start >= 0 and json_end > json_start:
                content = content[json_start:json_end+1]
            
            # Parser le JSON
            data = json.loads(content)
            
            # Vérifier et initialiser les listes si nécessaire
            list_keys = ["competences", "logiciels", "soft_skills", "experiences", "formation", "certifications"]
            for key in list_keys:
                if key not in data or not isinstance(data[key], list):
                    data[key] = []
            
            # S'assurer que les champs texte obligatoires existent
            text_keys = ["nom", "prenom", "poste", "adresse", "email", "telephone"]
            for key in text_keys:
                if key not in data or not isinstance(data[key], str):
                    data[key] = ""
            
            # Si des expériences sont présentes, effectuer un deuxième appel pour les raffiner
            if data.get("experiences") and len(data["experiences"]) > 0:
                data["experiences"] = refine_experiences(cv_text, data["experiences"])
            
            return data
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse JSON from OpenAI response: {content}")
            raise JSONParsingError(f"Invalid JSON format: {str(e)}")
            
    except openai.error.OpenAIError as e:
        logger.error(f"OpenAI API error: {str(e)}")
        raise OpenAIError(f"OpenAI API error: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error in OpenAI call: {str(e)}")
        raise OpenAIError(f"Unexpected error: {str(e)}")

@retry(
    stop=stop_after_attempt(2),
    wait=wait_exponential(multiplier=1, min=1, max=3),
    retry=retry_if_exception_type(OpenAIError),
    reraise=True
)
def refine_experiences(cv_text: str, initial_experiences: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Raffine les expériences professionnelles avec un second appel à l'API."""
    if not initial_experiences:
        return []
    
    try:
        # Limiter le contexte pour ne pas surcharger l'API
        cv_excerpt = cv_text[:8000] if len(cv_text) > 8000 else cv_text
        
        start_time = time.time()
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": EXPERIENCE_REFINEMENT_PROMPT},
                {"role": "user", "content": f"CV EXCERPT:\n{cv_excerpt}\n\nINITIAL EXPERIENCES:\n{json.dumps(initial_experiences, ensure_ascii=False, indent=2)}"}
            ],
            temperature=0.1,
            max_tokens=2000
        )
        
        execution_time = time.time() - start_time
        logger.info(f"Experience refinement completed in {execution_time:.2f} seconds")
        
        # Extraire et nettoyer la réponse
        content = response.choices[0].message.content.strip()
        
        # Nettoyer le contenu
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()
        
        # Extraire la partie JSON si nécessaire
        json_start = content.find('[')
        json_end = content.rfind(']')
        
        if json_start >= 0 and json_end > json_start:
            content = content[json_start:json_end+1]
        
        try:
            refined_experiences = json.loads(content)
            
            # Valider et normaliser les expériences raffinées
            if isinstance(refined_experiences, list):
                # Valider chaque expérience
                for exp in refined_experiences:
                    # Assurer que les clés requises sont présentes
                    required_keys = ["titre", "entreprise", "date_debut", "date_fin", "description"]
                    for key in required_keys:
                        if key not in exp:
                            exp[key] = ""
                    
                    # Assurer que les listes sont présentes
                    list_keys = ["responsabilites", "realisations", "technologies"]
                    for key in list_keys:
                        if key not in exp or not isinstance(exp[key], list):
                            exp[key] = []
                
                return refined_experiences
            else:
                logger.warning("Refined experiences is not a list, returning original")
                return initial_experiences
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse JSON from refinement response: {content}")
            return initial_experiences
            
    except openai.error.OpenAIError as e:
        logger.error(f"OpenAI API error during experience refinement: {str(e)}")
        return initial_experiences
    except Exception as e:
        logger.error(f"Unexpected error in experience refinement: {str(e)}")
        return initial_experiences

def normalize_date(date_str: str) -> str:
    """Normalise une date au format YYYY-MM."""
    if not date_str or date_str.lower() == "present":
        return date_str
        
    import re
    from datetime import datetime
    
    # Si déjà au format YYYY-MM
    if re.match(r'^\d{4}-\d{2}$', date_str):
        return date_str
        
    # Dictionnaire de correspondance des mois
    months_map = {
        'janvier': '01', 'january': '01', 'jan': '01',
        'février': '02', 'february': '02', 'feb': '02', 'fev': '02',
        'mars': '03', 'march': '03', 'mar': '03',
        'avril': '04', 'april': '04', 'apr': '04', 'avr': '04',
        'mai': '05', 'may': '05',
        'juin': '06', 'june': '06', 'jun': '06',
        'juillet': '07', 'july': '07', 'jul': '07', 'juil': '07',
        'août': '08', 'august': '08', 'aug': '08', 'aout': '08',
        'septembre': '09', 'september': '09', 'sep': '09', 'sept': '09',
        'octobre': '10', 'october': '10', 'oct': '10',
        'novembre': '11', 'november': '11', 'nov': '11',
        'décembre': '12', 'december': '12', 'dec': '12', 'déc': '12'
    }
    
    # Patterns pour différents formats de date
    patterns = [
        # YYYY-MM-DD
        (r'(\d{4})-(\d{1,2})-\d{1,2}', lambda m: f"{m.group(1)}-{int(m.group(2)):02d}"),
        # MM/YYYY or MM.YYYY
        (r'(\d{1,2})[/.] *(\d{4})', lambda m: f"{m.group(2)}-{int(m.group(1)):02d}"),
        # YYYY/MM or YYYY.MM
        (r'(\d{4})[/.] *(\d{1,2})', lambda m: f"{m.group(1)}-{int(m.group(2)):02d}"),
        # Month YYYY (ex: "Janvier 2020")
        (r'([a-zéû]+) +(\d{4})', lambda m: f"{m.group(2)}-{months_map.get(m.group(1).lower(), '01')}"),
        # YYYY Month (ex: "2020 Janvier")
        (r'(\d{4}) +([a-zéû]+)', lambda m: f"{m.group(1)}-{months_map.get(m.group(2).lower(), '01')}"),
        # Just YYYY
        (r'^(\d{4})$', lambda m: f"{m.group(1)}-01")
    ]
    
    for pattern, formatter in patterns:
        match = re.search(pattern, date_str, re.IGNORECASE)
        if match:
            try:
                return formatter(match)
            except (KeyError, IndexError):
                continue
    
    # Essai avec datetime
    try:
        for fmt in ['%d/%m/%Y', '%m/%d/%Y', '%Y-%m-%d', '%d-%m-%Y', '%m-%d-%Y', '%d.%m.%Y', '%m.%d.%Y', '%Y.%m.%d']:
            try:
                dt = datetime.strptime(date_str, fmt)
                return f"{dt.year}-{dt.month:02d}"
            except ValueError:
                continue
    except Exception:
        pass
        
    # Si rien ne fonctionne, retourner une valeur par défaut
    return ""

def fallback_extraction(text: str) -> Dict[str, Any]:
    """
    Méthode de secours pour extraire des informations basiques si OpenAI échoue
    Utilise des règles simples basées sur des regex pour obtenir au moins quelques données
    
    Args:
        text: Texte du CV
        
    Returns:
        Un dictionnaire avec les informations de base
    """
    import re
    
    # Structure par défaut
    result = {
        "nom": "",
        "prenom": "",
        "poste": "",
        "adresse": "",
        "competences": [],
        "logiciels": [],
        "soft_skills": [],
        "email": "",
        "telephone": "",
        "experiences": [],
        "formation": [],
        "certifications": []
    }
    
    # Extraction d'email (recherche basique)
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    if email_match:
        result["email"] = email_match.group(0)
    
    # Extraction de numéro de téléphone (format français)
    phone_matches = re.findall(r'(?:(?:\+|00)33|0)\s*[1-9](?:[\s.-]*\d{2}){4}', text)
    if phone_matches:
        result["telephone"] = phone_matches[0]
    
    # Extraction de compétences basiques (recherche de mots-clés courants)
    tech_keywords = ['python', 'java', 'javascript', 'html', 'css', 'sql', 'php', 
                    'react', 'angular', 'vue', 'node', 'django', 'flask', 'spring',
                    'c++', 'c#', 'docker', 'kubernetes', 'git', 'aws', 'azure']
    
    software_keywords = ['office', 'excel', 'word', 'powerpoint', 'photoshop', 
                       'illustrator', 'figma', 'sketch', 'jira', 'confluence',
                       'slack', 'trello', 'notion', 'adobe', 'autocad']
    
    soft_skills_keywords = ['communication', 'travail en équipe', 'leadership', 'autonomie',
                           'adaptabilité', 'gestion de projet', 'organisation', 'proactif',
                           'analyse', 'résolution de problèmes', 'créativité']
    
    for keyword in tech_keywords:
        if re.search(r'\b' + keyword + r'\b', text.lower()):
            result["competences"].append(keyword)
    
    for keyword in software_keywords:
        if re.search(r'\b' + keyword + r'\b', text.lower()):
            result["logiciels"].append(keyword)
            
    for keyword in soft_skills_keywords:
        if keyword in text.lower():
            result["soft_skills"].append(keyword)
    
    # Extraction basique d'expériences
    experience_sections = re.split(r'\n(?:expériences?|expériences? professionnelles?|parcours professionnel|professional experience|work experience)', text, flags=re.IGNORECASE)
    if len(experience_sections) > 1:
        exp_section = experience_sections[1].split('\n\n')[0]
        
        # Chercher des patterns d'emploi
        job_patterns = re.findall(r'(\d{4}[-/.]?\d{0,4})\s*[-–à]\s*(\d{4}[-/.]?\d{0,4}|[Aa]ctuel|[Pp]résent|[Pp]resent|[Cc]urrent)\s*[:\n]?\s*([^\n]*)', exp_section)
        
        for i, (start_date, end_date, job_desc) in enumerate(job_patterns):
            if i >= 3:  # Limiter à 3 expériences
                break
                
            # Essayer d'extraire l'entreprise et le poste
            job_parts = job_desc.split('|')
            title = ""
            company = ""
            
            if len(job_parts) >= 2:
                title = job_parts[0].strip()
                company = job_parts[1].strip()
            elif len(job_parts) == 1:
                title_company = job_parts[0].split('-', 1)
                if len(title_company) >= 2:
                    title = title_company[0].strip()
                    company = title_company[1].strip()
                else:
                    title = job_desc.strip()
            
            # Normaliser les dates
            start_date_normalized = normalize_date(start_date)
            end_date_normalized = "Present" if any(x in end_date.lower() for x in ['actuel', 'présent', 'present', 'current']) else normalize_date(end_date)
            
            result["experiences"].append({
                "titre": title,
                "entreprise": company,
                "lieu": "",
                "date_debut": start_date_normalized,
                "date_fin": end_date_normalized,
                "description": job_desc,
                "responsabilites": [],
                "realisations": [],
                "technologies": [],
                "taille_equipe": "",
                "type_contrat": ""
            })
    
    return result

def parse_cv(file: BinaryIO, filename: str, force_refresh: bool = False) -> Tuple[CVModel, bool]:
    """
    Fonction principale qui parse un CV en utilisant l'API OpenAI avec cache Redis.
    
    Args:
        file: Fichier CV (PDF ou DOCX)
        filename: Nom du fichier avec extension
        force_refresh: Forcer le rafraîchissement du cache
        
    Returns:
        Une tuple (CVModel, from_cache) avec les informations extraites et si elles viennent du cache
        
    Raises:
        CVParserError: Si le parsing du CV échoue
    """
    # Déterminer l'extension du fichier
    _, file_extension = os.path.splitext(filename)
    
    if file_extension.lower() not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Format de fichier non supporté: {file_extension}")
    
    # Calculer le hash du fichier pour la clé de cache
    file_hash = calculate_file_hash(file)
    cache_key = f"{CACHE_PREFIX}{file_hash}"
    
    # Vérifier si le résultat est déjà en cache et si on ne force pas le rafraîchissement
    if not force_refresh:
        cached_result = get_cached_result(cache_key)
        if cached_result:
            try:
                logger.info(f"Cache hit for file {filename}")
                # Créer une instance de CVModel avec les données du cache
                cv_model = CVModel(
                    nom=cached_result.get("nom", ""),
                    prenom=cached_result.get("prenom", ""),
                    poste=cached_result.get("poste", ""),
                    competences=cached_result.get("competences", []),
                    logiciels=cached_result.get("logiciels", []),
                    soft_skills=cached_result.get("soft_skills", []),
                    email=cached_result.get("email", ""),
                    telephone=cached_result.get("telephone", ""),
                    adresse=cached_result.get("adresse", "")
                )
                return cv_model, True
            except Exception as e:
                logger.warning(f"Error using cached data: {str(e)}. Will reprocess file.")
    
    # Si pas en cache ou erreur de cache, procéder au traitement
    try:
        # 1. Extraire le texte brut du document
        cv_text = extract_text(file, file_extension)
        
        if not cv_text.strip():
            logger.warning(f"Extracted empty text from file {filename}")
            raise FileExtractionError("Le texte extrait du fichier est vide")
        
        # 2. Tenter d'envoyer le texte à l'API OpenAI
        try:
            cv_data = parse_cv_with_openai(cv_text)
        except (OpenAIError, JSONParsingError) as e:
            logger.warning(f"OpenAI extraction failed, using fallback: {str(e)}")
            cv_data = fallback_extraction(cv_text)
        
        # 3. Mettre en cache le résultat
        cache_result(cache_key, cv_data)
        
        # 4. Créer une instance de CVModel avec les données extraites
        cv_model = CVModel(
            nom=cv_data.get("nom", ""),
            prenom=cv_data.get("prenom", ""),
            poste=cv_data.get("poste", ""),
            competences=cv_data.get("competences", []),
            logiciels=cv_data.get("logiciels", []),
            soft_skills=cv_data.get("soft_skills", []),
            email=cv_data.get("email", ""),
            telephone=cv_data.get("telephone", ""),
            adresse=cv_data.get("adresse", "")
        )
        
        return cv_model, False
        
    except FileExtractionError as e:
        logger.error(f"File extraction error: {str(e)}")
        raise CVParserError(f"Erreur d'extraction du fichier: {str(e)}")
    except Exception as e:
        logger.error(f"Error parsing CV: {str(e)}", exc_info=True)
        raise CVParserError(f"Erreur lors du parsing du CV: {str(e)}")
