
# CV Parser Service - Service de parsing CV

import os
import time
import logging
from typing import Dict, Any, Optional, List, BinaryIO
import tempfile
import json
import traceback
import re

from app.core.config import settings
from app.services.resilience import resilient_openai_call
from app.services.mock_parser import get_mock_cv_data

# Setup logging
logger = logging.getLogger(__name__)

def parse_cv(file_path: str, file_format: Optional[str] = None) -> Dict[str, Any]:
    """Parse un CV pour en extraire les informations structurées
    
    Args:
        file_path: Chemin vers le fichier CV
        file_format: Format du fichier (.pdf, .docx, etc.)
        
    Returns:
        Dict[str, Any]: Informations structurées extraites du CV
    """
    # 1. Déterminer le format si non fourni
    if not file_format and file_path:
        file_format = os.path.splitext(file_path)[1].lower()
    
    # Logging du fichier traité pour le debugging
    logger.info(f"Traitement du fichier: {os.path.basename(file_path)} (format: {file_format})")
    
    try:
        # 2. Extraire le texte du CV selon le format
        cv_text = extract_text_from_file(file_path, file_format)
        
        # Log de la taille du texte extrait pour debug
        logger.info(f"Texte extrait: {len(cv_text)} caractères")
        if len(cv_text) < 100:  # Si le texte extrait est très court, c'est probablement un problème
            logger.warning(f"Texte extrait très court ({len(cv_text)} caractères), possible problème d'extraction")
            logger.debug(f"Contenu extrait: {cv_text}")
        
        # Pré-traitement du texte pour améliorer la détection
        cv_text = preprocess_cv_text(cv_text)
        
        # 3. Utiliser OpenAI pour analyser le CV
        start_time = time.time()
        
        try:
            # Si USE_MOCK_PARSER est activé, utiliser le mock au lieu de l'API
            if settings.USE_MOCK_PARSER:
                logger.info(f"Utilisation du mock parser (mode de simulation) pour {file_path}")
                parsed_data = get_mock_cv_data(cv_text, os.path.basename(file_path))
            else:
                # Sinon, utiliser l'API OpenAI
                parsed_data = analyze_cv_with_gpt(cv_text)
                
                # Post-traitement pour corriger et enrichir les données
                parsed_data = postprocess_cv_data(parsed_data, cv_text)
        except Exception as e:
            logger.error(f"Erreur lors de l'analyse du CV: {str(e)}. Fallback sur le mock parser.")
            logger.error(f"Stacktrace: {traceback.format_exc()}")
            # En cas d'erreur, utiliser le mock parser comme fallback
            parsed_data = get_mock_cv_data(cv_text, os.path.basename(file_path))
        
        processing_time = time.time() - start_time
        logger.info(f"CV parsé en {processing_time:.2f} secondes")
        
        # 4. Ajouter des métadonnées au résultat
        result = {
            "processing_time": processing_time,
            "parsed_at": time.time(),
            "file_format": file_format,
            "model": "mock" if settings.USE_MOCK_PARSER else settings.OPENAI_MODEL,
            "data": parsed_data
        }
        
        return result
    except Exception as e:
        logger.error(f"Erreur pendant le parsing du CV {os.path.basename(file_path)}: {str(e)}")
        logger.error(f"Stacktrace: {traceback.format_exc()}")
        
        # Retourner un résultat avec l'erreur mais une structure minimale valide
        return {
            "processing_time": 0,
            "parsed_at": time.time(),
            "file_format": file_format,
            "model": "error",
            "error": str(e),
            "data": {
                "personal_info": {"name": "", "email": "", "phone": ""},
                "position": "",
                "skills": [],
                "experience": [],
                "education": [],
                "languages": [],
                "softwares": []
            }
        }

def preprocess_cv_text(text: str) -> str:
    """Prétraite le texte du CV pour améliorer la qualité de l'analyse"""
    if not text:
        return ""
        
    # Remplacer les séquences de plusieurs espaces par un seul
    text = re.sub(r' +', ' ', text)
    
    # Remplacer les séquences de plusieurs lignes vides par une seule
    text = re.sub(r'\n+', '\n', text)
    
    # Essayer de détecter et de reconstruire les numéros de téléphone fragmentés
    # Exemple: +33 6 12 34 56 78 --> +33612345678
    phone_patterns = [
        r'(\+\d{2})\s*(\d)\s*(\d{2})\s*(\d{2})\s*(\d{2})\s*(\d{2})',  # +33 6 12 34 56 78
        r'0\s*(\d)\s*(\d{2})\s*(\d{2})\s*(\d{2})\s*(\d{2})',          # 06 12 34 56 78
    ]
    
    for pattern in phone_patterns:
        text = re.sub(pattern, lambda m: ''.join(m.groups()) if len(m.groups()) > 1 else m.group(0), text)
    
    return text

def postprocess_cv_data(data: Dict[str, Any], original_text: str) -> Dict[str, Any]:
    """Post-traitement des données extraites pour corriger et enrichir les informations"""
    if not data:
        return {}
    
    # Correction du champ name pour enlever "undefined" si présent
    if "personal_info" in data and "name" in data["personal_info"]:
        # Enlève le préfixe "undefined" des noms
        if data["personal_info"]["name"] and data["personal_info"]["name"].startswith("undefined "):
            data["personal_info"]["name"] = data["personal_info"]["name"].replace("undefined ", "")
    
    # Si le téléphone n'est pas détecté, essayer de le trouver dans le texte
    if ("personal_info" not in data or 
        "phone" not in data["personal_info"] or 
        not data["personal_info"]["phone"]):
        phone = extract_phone_from_text(original_text)
        if phone:
            if "personal_info" not in data:
                data["personal_info"] = {}
            data["personal_info"]["phone"] = phone
    
    # Extraire le titre de poste s'il n'est pas détecté
    if not data.get("position"):
        position = extract_position_from_text(original_text)
        if position:
            data["position"] = position
    
    # Assurez-vous que les listes clés existent
    for key in ["skills", "languages", "softwares", "experience", "education"]:
        if key not in data:
            data[key] = []
    
    # Traiter les compétences et logiciels si peu ont été détectés
    if len(data.get("skills", [])) < 3 or len(data.get("softwares", [])) < 1:
        skills, softwares = extract_skills_and_software(original_text)
        # Fusionner avec les compétences existantes
        if skills:
            existing_skills = set(s.get("name", s) if isinstance(s, dict) else s for s in data.get("skills", []))
            for skill in skills:
                if skill not in existing_skills:
                    data["skills"].append({"name": skill})
        
        # Fusionner avec les logiciels existants
        if softwares:
            existing_softwares = set(data.get("softwares", []))
            for software in softwares:
                if software not in existing_softwares:
                    data["softwares"].append(software)
    
    # Extraire les langues si aucune n'a été détectée
    if not data.get("languages"):
        languages = extract_languages(original_text)
        if languages:
            data["languages"] = languages
    
    return data

def extract_phone_from_text(text: str) -> str:
    """Extrait un numéro de téléphone du texte avec divers formats"""
    phone_patterns = [
        r'\+[\d\s]{10,15}',                            # +33 6 12 34 56 78
        r'0\d[\s.-]?\d{2}[\s.-]?\d{2}[\s.-]?\d{2}[\s.-]?\d{2}',  # 06 12 34 56 78
        r'\d{3}[\s.-]?\d{3}[\s.-]?\d{4}',              # 123-456-7890
        r'\(\d{3}\)[\s.-]?\d{3}[\s.-]?\d{4}',          # (123) 456-7890
        r'\d{2}[\s.-]?\d{2}[\s.-]?\d{2}[\s.-]?\d{2}[\s.-]?\d{2}'  # 01 23 45 67 89
    ]
    
    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        if matches:
            # Nettoyer le numéro trouvé (enlever espaces, tirets, etc)
            phone = re.sub(r'[\s.-]', '', matches[0])
            # Formater joliment si c'est un numéro français
            if phone.startswith('0') and len(phone) == 10:
                return f"{phone[0:2]} {phone[2:4]} {phone[4:6]} {phone[6:8]} {phone[8:10]}"
            return phone
    
    return ""

def extract_position_from_text(text: str) -> str:
    """Extrait le titre de poste du texte"""
    # Recherche de motifs courants pour les titres de postes
    position_patterns = [
        r'(?i)poste actuel\s*:?\s*([^\n]+)',
        r'(?i)poste recherché\s*:?\s*([^\n]+)',
        r'(?i)intitulé du poste\s*:?\s*([^\n]+)',
        r'(?i)titre\s*:?\s*([^\n]+)',
        r'(?i)(comptable|auditeur|contrôleur|gestionnaire|analyste|manager|directeur)[^\n]{0,50}',
        r'(?i)(consultant|expert|responsable|chargé)[^\n]{0,50}',
    ]
    
    for pattern in position_patterns:
        matches = re.findall(pattern, text)
        if matches:
            # Limiter la longueur du titre
            position = matches[0].strip()
            if len(position) > 5 and len(position) < 100:  # Filtre les résultats trop courts ou trop longs
                return position
    
    # Recherche basée sur les premières lignes du CV (souvent titre)
    first_lines = text.split('\n')[:10]
    for line in first_lines:
        line = line.strip()
        if len(line) > 5 and len(line) < 50:
            if any(keyword in line.lower() for keyword in [
                'comptable', 'auditeur', 'contrôleur', 'gestionnaire', 
                'analyste', 'manager', 'directeur', 'consultant', 
                'expert', 'responsable', 'chargé'
            ]):
                return line
    
    return ""

def extract_skills_and_software(text: str) -> tuple:
    """Extrait des compétences et logiciels du texte du CV"""
    skills = []
    softwares = []
    
    # Liste de compétences courantes en finance/comptabilité
    accounting_skills = [
        'Comptabilité générale', 'Comptabilité analytique', 'Comptabilité clients', 
        'Comptabilité fournisseurs', 'Fiscalité', 'Audit', 'Contrôle de gestion',
        'Gestion de trésorerie', 'Finance d\'entreprise', 'Normes IFRS', 'Normes US GAAP',
        'Consolidation', 'Reporting', 'Budget', 'Prévisions', 'Analyse financière',
        'Clôture comptable', 'Rapprochement bancaire', 'Liasse fiscale', 'Bilan'
    ]
    
    # Liste de logiciels courants
    common_software = [
        'SAP', 'Oracle', 'Sage', 'Cegid', 'Microsoft Office', 'Excel', 'Word', 'PowerPoint',
        'Access', 'Dynamics', 'ADP', 'QuickBooks', 'Microsoft Dynamics', 'EBP', 'Coala',
        'Talentia', 'AX', 'Navision', 'JD Edwards', 'PeopleSoft', 'Workday', 'Xero',
        'MYOB', 'Freshbooks', 'NetSuite', 'Odoo', 'Wave', 'Zoho Books', 'SQL', 'Power BI',
        'Tableau', 'Quadra', 'Ciel', 'FEC-Expert', 'Exact', 'SalesForce', 'AX 365'
    ]
    
    # Vérifier les compétences comptables
    for skill in accounting_skills:
        if re.search(r'\b' + re.escape(skill) + r'\b', text, re.IGNORECASE):
            skills.append(skill)
    
    # Vérifier les logiciels
    for software in common_software:
        if re.search(r'\b' + re.escape(software) + r'\b', text, re.IGNORECASE):
            softwares.append(software)
    
    return skills, softwares

def extract_languages(text: str) -> List[Dict[str, str]]:
    """Extrait les langues et niveaux du texte du CV"""
    languages = []
    
    # Liste des langues courantes
    common_languages = [
        'Français', 'Anglais', 'Espagnol', 'Allemand', 'Italien', 'Portugais',
        'Russe', 'Chinois', 'Arabe', 'Japonais', 'Néerlandais', 'Suédois'
    ]
    
    # Niveaux courants
    levels = [
        'natif', 'native', 'maternelle', 'courant', 'professionnel', 'bilingue',
        'intermédiaire', 'scolaire', 'notions', 'débutant', 'avancé', 'B1', 'B2', 'C1', 'C2', 'A1', 'A2'
    ]
    
    # Recherche de motifs comme "Anglais: courant"
    for language in common_languages:
        pattern = rf'(?i){re.escape(language)}[\s:,]+([^\n,.;]+)'
        matches = re.findall(pattern, text)
        
        if matches:
            level_found = False
            for level_text in matches:
                for level in levels:
                    if re.search(rf'\b{re.escape(level)}\b', level_text, re.IGNORECASE):
                        languages.append({
                            "language": language,
                            "level": level
                        })
                        level_found = True
                        break
                if level_found:
                    break
            
            # Si aucun niveau n'est trouvé mais la langue est mentionnée
            if not level_found:
                languages.append({
                    "language": language,
                    "level": "mentionné"
                })
        # Recherche simple de la présence de la langue
        elif re.search(rf'\b{re.escape(language)}\b', text, re.IGNORECASE):
            languages.append({
                "language": language,
                "level": "mentionné"
            })
    
    return languages

def extract_text_from_file(file_path: str, file_format: Optional[str] = None) -> str:
    """Extrait le texte d'un fichier CV
    
    Args:
        file_path: Chemin vers le fichier
        file_format: Format du fichier
        
    Returns:
        str: Texte extrait du CV
    """
    logger.info(f"Extraction du texte depuis {file_path} (format: {file_format})")
    
    # Déterminer le format si non fourni
    if not file_format:
        file_format = os.path.splitext(file_path)[1].lower()
    
    # Vérifier si le fichier existe pour éviter des erreurs
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Le fichier n'existe pas: {file_path}")
    
    # Log de la taille du fichier
    file_size = os.path.getsize(file_path)
    logger.info(f"Taille du fichier: {file_size / 1024:.2f} KB")
    
    # Extraction selon le format avec gestion d'erreur améliorée
    try:
        if file_format.lower() in [".pdf", ".PDF"]:
            return extract_text_from_pdf(file_path)
        elif file_format.lower() in [".docx", ".DOCX"]:
            return extract_text_from_docx(file_path)
        elif file_format.lower() in [".doc", ".DOC"]:
            return extract_text_from_doc(file_path)
        elif file_format.lower() in [".txt", ".TXT", ".text"]:
            return extract_text_from_txt(file_path)
        elif file_format.lower() in [".rtf", ".RTF"]:
            return extract_text_from_rtf(file_path)
        else:
            # Tenter une extraction générique pour les formats non reconnus
            logger.warning(f"Format non reconnu: {file_format}. Tentative d'extraction générique.")
            return extract_text_generic(file_path)
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte ({file_format}): {str(e)}")
        logger.error(f"Stacktrace: {traceback.format_exc()}")
        # En cas d'échec, tenter une extraction générique alternative
        try:
            logger.info("Tentative d'extraction alternative...")
            return extract_text_generic(file_path)
        except Exception as alt_e:
            logger.error(f"Échec de l'extraction alternative: {str(alt_e)}")
            # Retourner une chaîne vide mais non None pour éviter des erreurs en aval
            return f"Échec d'extraction du texte. Format: {file_format}. Erreur: {str(e)}"

def extract_text_generic(file_path: str) -> str:
    """Méthode d'extraction générique qui tente plusieurs approches"""
    logger.info(f"Extraction générique pour {file_path}")
    
    # Essayer des méthodes alternatives d'extraction
    try:
        # Tenter avec textract qui supporte de nombreux formats
        try:
            import textract
            text = textract.process(file_path).decode('utf-8')
            if text and len(text) > 100:  # Vérifier qu'on a extrait quelque chose de significatif
                return text
        except:
            logger.warning("Échec de l'extraction avec textract")
        
        # Tenter avec pdfplumber (autre bibliothèque pour PDF)
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
                if text:
                    return text
        except:
            logger.warning("Échec de l'extraction avec pdfplumber")
        
        # Utiliser notre propre méthode PDF comme dernier recours
        try:
            return extract_text_from_pdf(file_path)
        except:
            logger.warning("Échec de l'extraction PDF standard")
        
        # Si tout échoue, tenter une lecture binaire simple
        with open(file_path, 'rb') as f:
            content = f.read()
            try:
                # Tenter plusieurs encodages
                for encoding in ['utf-8', 'latin-1', 'windows-1252', 'ascii']:
                    try:
                        return content.decode(encoding)
                    except:
                        continue
            except:
                pass
        
        return "Extraction de texte échouée pour ce document."
    except Exception as e:
        logger.error(f"Toutes les méthodes d'extraction ont échoué: {str(e)}")
        return "Échec de toutes les méthodes d'extraction de texte."

def extract_text_from_pdf(file_path: str) -> str:
    """Extrait le texte d'un fichier PDF avec meilleure gestion d'erreurs"""
    try:
        logger.info(f"Tentative d'extraction PDF depuis {file_path}")
        
        # Première tentative avec PyPDF2
        try:
            from PyPDF2 import PdfReader
            
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                text = ""
                for page in reader.pages:
                    extracted_text = page.extract_text()
                    if extracted_text:
                        text += extracted_text + "\n"
                
                if text.strip():  # Vérifier que le texte extrait n'est pas vide
                    logger.info(f"Extraction PyPDF2 réussie: {len(text)} caractères")
                    return text
                else:
                    logger.warning("PyPDF2 n'a pas extrait de texte - le PDF pourrait être scanné ou contenir des images")
        except Exception as e:
            logger.warning(f"Échec de l'extraction avec PyPDF2: {str(e)}")
        
        # Deuxième tentative avec pdfminer.six
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            text = pdfminer_extract(file_path)
            if text.strip():
                logger.info(f"Extraction pdfminer.six réussie: {len(text)} caractères")
                return text
            else:
                logger.warning("pdfminer.six n'a pas extrait de texte")
        except Exception as e:
            logger.warning(f"Échec de l'extraction avec pdfminer.six: {str(e)}")
        
        # Troisième tentative avec OCR si les autres méthodes échouent
        try:
            import pytesseract
            from PIL import Image
            import pdf2image
            
            logger.info("Tentative d'extraction via OCR (conversion PDF en images puis OCR)")
            pages = pdf2image.convert_from_path(file_path)
            text = ""
            for page in pages:
                text += pytesseract.image_to_string(page) + "\n"
            
            if text.strip():
                logger.info(f"Extraction OCR réussie: {len(text)} caractères")
                return text
        except Exception as e:
            logger.warning(f"Échec de l'extraction avec OCR: {str(e)}")
        
        # Si toutes les méthodes échouent
        return "Texte non extractible de ce PDF. Il pourrait s'agir d'un document scanné ou protégé."
        
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte PDF: {str(e)}")
        raise

def extract_text_from_docx(file_path: str) -> str:
    """Extrait le texte d'un fichier DOCX"""
    try:
        import docx
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extraire également les tableaux qui peuvent contenir des informations importantes
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return text
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte DOCX: {str(e)}")
        raise

def extract_text_from_doc(file_path: str) -> str:
    """Extrait le texte d'un fichier DOC (ancien format Word)"""
    try:
        # Plusieurs approches possibles
        methods_tried = []
        
        # Essayer avec textract (nécessite installation système)
        try:
            import textract
            text = textract.process(file_path).decode('utf-8')
            methods_tried.append("textract")
            if text:
                return text
        except ImportError:
            logger.warning("Module textract non disponible pour l'extraction DOC")
        except Exception as e:
            logger.warning(f"Échec de l'extraction DOC avec textract: {str(e)}")
        
        # Essayer avec antiword
        try:
            import subprocess
            result = subprocess.run(['antiword', file_path], stdout=subprocess.PIPE)
            text = result.stdout.decode('utf-8')
            methods_tried.append("antiword")
            if text:
                return text
        except Exception as e:
            logger.warning(f"Échec de l'extraction DOC avec antiword: {str(e)}")
        
        # Essayer avec pywin32 (Windows uniquement)
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(file_path)
            text = doc.Content.Text
            doc.Close()
            word.Quit()
            methods_tried.append("win32com")
            if text:
                return text
        except Exception as e:
            logger.warning(f"Échec de l'extraction DOC avec win32com: {str(e)}")

        if not methods_tried:
            raise NotImplementedError("Aucune méthode d'extraction DOC n'a fonctionné")
        else:
            raise Exception(f"Échec de l'extraction DOC avec les méthodes: {', '.join(methods_tried)}")
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte DOC: {str(e)}")
        raise

def extract_text_from_txt(file_path: str) -> str:
    """Extrait le texte d'un fichier texte"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Essayer avec des encodages alternatifs
        for encoding in ['latin-1', 'windows-1252', 'iso-8859-1', 'cp1252']:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        # Si tous les encodages échouent, essayer en mode binaire
        with open(file_path, 'rb') as file:
            content = file.read()
            return str(content)
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte TXT: {str(e)}")
        raise

def extract_text_from_rtf(file_path: str) -> str:
    """Extrait le texte d'un fichier RTF"""
    try:
        # Essayer avec striprtf
        try:
            from striprtf.striprtf import rtf_to_text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            return rtf_to_text(content)
        except ImportError:
            # Fallback à textract
            import textract
            text = textract.process(file_path).decode('utf-8')
            return text
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte RTF: {str(e)}")
        raise

def analyze_cv_with_gpt(cv_text: str) -> Dict[str, Any]:
    """Analyse un CV avec GPT pour en extraire les informations structurées
    
    Args:
        cv_text: Texte du CV
        
    Returns:
        Dict[str, Any]: Informations structurées extraites du CV
    """
    logger.info(f"Analyse du CV avec {settings.OPENAI_MODEL} (longueur: {len(cv_text)} caractères)")
    
    # Si le texte est vide ou trop court, retourner une structure vide
    if not cv_text or len(cv_text) < 50:
        logger.warning(f"Texte CV trop court ({len(cv_text)} caractères), impossible d'analyser")
        return {
            "personal_info": {"name": "", "email": "", "phone": ""},
            "position": "",
            "skills": [],
            "experience": [],
            "education": [],
            "languages": [],
            "softwares": []
        }
    
    # Si le texte est trop long, le tronquer pour éviter des problèmes avec l'API
    max_tokens = 15000  # Approximativement 15000 caractères
    if len(cv_text) > max_tokens:
        logger.warning(f"CV trop long ({len(cv_text)} caractères), troncature à {max_tokens} caractères")
        cv_text = cv_text[:max_tokens] + "...[texte tronqué]"
    
    # Définir le prompt pour l'extraction d'information structurée avec instructions améliorées
    prompt = f"""
Tu es un expert en analyse de CV et extraction de données pour l'industrie du recrutement.
Tu dois extraire avec précision toutes les informations importantes d'un CV, en particulier pour les domaines de la finance, comptabilité et audit.

INSTRUCTIONS IMPÉRATIVES:
1. Extrait UNIQUEMENT les informations réellement présentes dans le CV.
2. Ne génère JAMAIS d'informations fictives comme "John Doe" ou "example@email.com".
3. Pour tout champ non présent dans le CV, renvoie une valeur vide (chaîne vide ou tableau vide).
4. Sois particulièrement attentif aux numéros de téléphone, titres de poste, langues et logiciels.
5. Pour les compétences, différencie bien les compétences techniques et les logiciels maîtrisés.

EXTRACTION DEMANDÉE:
Retourne un JSON avec cette structure précise et TOUS ces champs, même vides:

{{
  "personal_info": {{
    "name": "",     // Nom complet sans préfixe comme "undefined"
    "email": "",    // Email exact
    "phone": "",    // Numéro de téléphone dans son format original
    "address": ""   // Adresse si présente, sinon vide
  }},
  "position": "",   // Poste actuel ou recherché (titre professionnel)
  "skills": [       // Liste des compétences (hors langues et logiciels)
    {{
      "name": "Compétence 1"
    }},
    {{
      "name": "Compétence 2"
    }}
  ],
  "softwares": [    // Logiciels maîtrisés (SAP, Excel, Sage, etc.)
    "Logiciel 1",
    "Logiciel 2"
  ],
  "languages": [    // Langues
    {{
      "language": "Français",
      "level": "Natif"
    }},
    {{
      "language": "Anglais",
      "level": "Courant"
    }}
  ],
  "experience": [   // Expériences professionnelles
    {{
      "title": "Titre du poste",
      "company": "Nom de l'entreprise",
      "start_date": "Date de début",
      "end_date": "Date de fin ou Présent",
      "description": "Description des responsabilités"
    }}
  ],
  "education": [    // Formation
    {{
      "degree": "Nom du diplôme",
      "institution": "Établissement",
      "start_date": "Date de début",
      "end_date": "Date de fin"
    }}
  ]
}}

CV À ANALYSER:
{cv_text}

IMPORTANT: Tu dois retourner UNIQUEMENT le JSON avec toutes les informations récupérées du CV, sans aucun texte d'introduction ou commentaire.
"""
    
    try:
        # Appel résilient à OpenAI (avec circuit breaker et retry)
        response_text = resilient_openai_call(
            prompt=prompt, 
            model=settings.OPENAI_MODEL,
            temperature=0.1,
            max_tokens=4000
        )
        
        logger.debug(f"Réponse brute d'OpenAI: {response_text}")
        
        # Parser la réponse JSON
        try:
            # Nettoyage du texte pour s'assurer qu'il ne contient que du JSON
            import re
            json_pattern = r'(\{[\s\S]*\})' 
            match = re.search(json_pattern, response_text)
            if match:
                json_str = match.group(1)
                try:
                    parsed_result = json.loads(json_str)
                    logger.info("Parsing JSON réussi")
                    return parsed_result
                except json.JSONDecodeError as json_err:
                    logger.error(f"Erreur de décodage JSON après extraction: {str(json_err)}")
            
            # Essayer de parser directement si l'extraction a échoué
            try:
                parsed_result = json.loads(response_text)
                return parsed_result
            except json.JSONDecodeError:
                pass
                
            # Tentative de nettoyage et d'extraction plus agressive
            clean_text = response_text.strip()
            if clean_text.startswith("```json"):
                clean_text = clean_text[7:]
            if clean_text.endswith("```"):
                clean_text = clean_text[:-3]
            clean_text = clean_text.strip()
            
            try:
                return json.loads(clean_text)
            except json.JSONDecodeError:
                logger.error("Échec de la tentative de nettoyage JSON")
                
                # En dernier recours, renvoyer un dictionnaire avec la réponse brute
                return {
                    "error": "Format JSON invalide dans la réponse",
                    "raw_response": response_text[:1000],  # Tronquer pour éviter les réponses trop longues
                    "personal_info": {
                        "name": "",
                        "email": "",
                        "phone": ""
                    },
                    "position": "",
                    "skills": [],
                    "experience": [],
                    "education": [],
                    "languages": [],
                    "softwares": []
                }
        except Exception as parse_err:
            logger.error(f"Erreur lors du parsing de la réponse: {str(parse_err)}")
            # Structure de retour par défaut en cas d'erreur
            return {
                "error": f"Erreur de parsing: {str(parse_err)}",
                "personal_info": {"name": "", "email": "", "phone": ""},
                "position": "",
                "skills": [],
                "experience": [],
                "education": [],
                "languages": [],
                "softwares": []
            }
    
    except Exception as e:
        logger.error(f"Erreur lors de l'analyse du CV avec GPT: {str(e)}")
        # Structure de retour par défaut en cas d'erreur
        return {
            "error": f"Erreur d'analyse: {str(e)}",
            "personal_info": {"name": "", "email": "", "phone": ""},
            "position": "",
            "skills": [],
            "experience": [],
            "education": [],
            "languages": [],
            "softwares": []
        }
