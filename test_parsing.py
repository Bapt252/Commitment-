#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
SuperSmartMatch V3.0 Enhanced - Testeur de Parsing Avancé
Diagnostic complet du système de parsing de CV et PDF
"""

import os
import sys
import time
import logging
from pathlib import Path
import tempfile
from typing import Dict, List, Optional, Any

# Imports pour parsing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️ PyPDF2 non disponible")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx non disponible")

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️ pytesseract/Pillow non disponible")

import re
import json
from datetime import datetime

# Configuration logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ParseTester:
    """Testeur avancé du système de parsing"""
    
    def __init__(self):
        self.results = []
        self.stats = {
            'total_files': 0,
            'successful_extractions': 0,
            'failed_extractions': 0,
            'empty_extractions': 0,
            'pdf_count': 0,
            'docx_count': 0,
            'image_count': 0,
            'txt_count': 0
        }
        
        # Base de compétences (version réduite pour test)
        self.skills_database = {
            "tech": [
                "Python", "Java", "JavaScript", "TypeScript", "React", "Vue.js", "Angular",
                "Node.js", "Django", "Flask", "Spring", "Docker", "Kubernetes", "AWS",
                "Azure", "GCP", "DevOps", "CI/CD", "Git", "SQL", "PostgreSQL", "MongoDB",
                "Redis", "Machine Learning", "AI", "Data Science", "TensorFlow", "PyTorch",
                "API", "REST", "GraphQL", "Microservices", "Linux", "Bash", "PowerShell"
            ],
            "juridique": [
                "Droit", "Juridique", "Contentieux", "Contrats", "RGPD", "Compliance",
                "Administrative", "Conseil juridique", "Rédaction juridique", "Veille juridique",
                "Droit social", "Droit commercial", "Droit des affaires", "Propriété intellectuelle"
            ],
            "rh": [
                "Ressources Humaines", "RH", "Recrutement", "Formation", "Paie", "Gestion RH",
                "Talent Management", "Performance Management", "SIRH", "Relations sociales"
            ],
            "business": [
                "Management", "Leadership", "Strategy", "Business Development", "Marketing",
                "Sales", "Finance", "Accounting", "Budget", "Controlling", "Analytics",
                "Project Management", "Change Management", "Innovation"
            ],
            "langues": [
                "Français", "Anglais", "Espagnol", "Allemand", "Italien", "Chinois",
                "Japonais", "Arabe", "Russe", "Portugais"
            ]
        }
    
    def print_header(self):
        """Header de test"""
        print("\n" + "="*70)
        print("🔍 SUPERSMARTMATCH V3.0 - TESTEUR DE PARSING AVANCÉ")
        print("="*70)
        print(f"📅 Test démarré: {datetime.now().strftime('%H:%M:%S')}")
        print("\n🔧 Vérification des dépendances:")
        print(f"   📄 PyPDF2: {'✅' if PDF_AVAILABLE else '❌'}")
        print(f"   📝 python-docx: {'✅' if DOCX_AVAILABLE else '❌'}")
        print(f"   🖼️ OCR (pytesseract): {'✅' if OCR_AVAILABLE else '❌'}")
        print("="*70)
    
    def test_text_extraction(self, file_path: str) -> Dict[str, Any]:
        """Test extraction de texte selon le format"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {"error": f"Fichier non trouvé: {file_path}"}
        
        result = {
            "file": str(file_path),
            "size": file_path.stat().st_size,
            "extension": file_path.suffix.lower(),
            "text": "",
            "error": None,
            "processing_time": 0,
            "method": "unknown"
        }
        
        start_time = time.time()
        
        try:
            ext = file_path.suffix.lower()
            
            if ext == '.pdf':
                result["text"] = self._extract_pdf(file_path)
                result["method"] = "PyPDF2"
                self.stats['pdf_count'] += 1
                
            elif ext in ['.docx', '.doc']:
                result["text"] = self._extract_docx(file_path)
                result["method"] = "python-docx"
                self.stats['docx_count'] += 1
                
            elif ext in ['.png', '.jpg', '.jpeg']:
                result["text"] = self._extract_image(file_path)
                result["method"] = "pytesseract OCR"
                self.stats['image_count'] += 1
                
            elif ext == '.txt':
                result["text"] = self._extract_txt(file_path)
                result["method"] = "lecture directe"
                self.stats['txt_count'] += 1
                
            else:
                result["error"] = f"Format non supporté: {ext}"
            
            result["processing_time"] = (time.time() - start_time) * 1000
            
            # Classification du résultat
            if result["error"]:
                self.stats['failed_extractions'] += 1
            elif len(result["text"].strip()) == 0:
                self.stats['empty_extractions'] += 1
            else:
                self.stats['successful_extractions'] += 1
            
            self.stats['total_files'] += 1
            
        except Exception as e:
            result["error"] = str(e)
            result["processing_time"] = (time.time() - start_time) * 1000
            self.stats['failed_extractions'] += 1
            self.stats['total_files'] += 1
        
        return result
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extraction PDF avec PyPDF2"""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 non disponible")
        
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            print(f"   📄 PDF détecté: {len(pdf_reader.pages)} page(s)")
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    text += page_text
                    print(f"      Page {page_num + 1}: {len(page_text)} caractères extraits")
                except Exception as e:
                    print(f"      Page {page_num + 1}: Erreur - {e}")
        
        return text
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extraction DOCX avec python-docx"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx non disponible")
        
        doc = docx.Document(file_path)
        text = ""
        
        print(f"   📝 DOCX détecté: {len(doc.paragraphs)} paragraphe(s)")
        
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        return text
    
    def _extract_image(self, file_path: Path) -> str:
        """Extraction image avec OCR"""
        if not OCR_AVAILABLE:
            raise ImportError("pytesseract/Pillow non disponible")
        
        print(f"   🖼️ Image détectée: OCR en cours...")
        
        image = Image.open(file_path)
        
        # Informations sur l'image
        print(f"      Dimensions: {image.size}")
        print(f"      Mode: {image.mode}")
        
        # OCR avec langues française et anglaise
        text = pytesseract.image_to_string(image, lang='fra+eng')
        
        print(f"      OCR terminé: {len(text)} caractères extraits")
        
        return text
    
    def _extract_txt(self, file_path: Path) -> str:
        """Extraction TXT"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def analyze_extracted_data(self, text: str) -> Dict[str, Any]:
        """Analyse du texte extrait pour détecter les données CV"""
        analysis = {
            "text_length": len(text),
            "lines_count": len(text.split('\n')),
            "words_count": len(text.split()),
            "name": self._extract_name(text),
            "skills": self._extract_skills(text),
            "experience_years": self._extract_experience(text),
            "education": self._extract_education(text),
            "languages": self._extract_languages(text),
            "sector": None,
            "emails": self._extract_emails(text),
            "phones": self._extract_phones(text)
        }
        
        # Détection secteur basé sur compétences
        analysis["sector"] = self._detect_sector(analysis["skills"])
        
        return analysis
    
    def _extract_name(self, text: str) -> Optional[str]:
        """Extraction nom candidat"""
        lines = text.split('\n')[:5]
        
        for line in lines:
            line = line.strip()
            # Pattern nom/prénom simple
            if re.match(r'^[A-ZÀ-Ÿ][a-zà-ÿ]+\s+[A-ZÀ-Ÿ][a-zà-ÿ]+', line):
                return line
        
        return None
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extraction compétences"""
        skills_found = []
        text_lower = text.lower()
        
        for sector, skills_list in self.skills_database.items():
            for skill in skills_list:
                if skill.lower() in text_lower:
                    skills_found.append(skill)
        
        return list(set(skills_found))
    
    def _extract_experience(self, text: str) -> int:
        """Extraction années d'expérience"""
        patterns = [
            r'(\d+)\s*ans?\s*d.expérience',
            r'(\d+)\s*années?\s*d.expérience',
            r'(\d+)\s*ans?\s*dans',
            r'expérience.*?(\d+)\s*ans?'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text.lower())
            if matches:
                return max(int(match) for match in matches)
        
        return 0
    
    def _extract_education(self, text: str) -> List[str]:
        """Extraction formation"""
        education_keywords = [
            'Master', 'Licence', 'Bachelor', 'BTS', 'DUT', 'MBA',
            'Doctorat', 'PhD', 'Ingénieur', 'École', 'Université'
        ]
        
        found = []
        for keyword in education_keywords:
            if keyword.lower() in text.lower():
                found.append(keyword)
        
        return found
    
    def _extract_languages(self, text: str) -> List[str]:
        """Extraction langues"""
        languages_found = []
        text_lower = text.lower()
        
        for language in self.skills_database["langues"]:
            if language.lower() in text_lower:
                languages_found.append(language)
        
        return languages_found
    
    def _extract_emails(self, text: str) -> List[str]:
        """Extraction emails"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        return re.findall(email_pattern, text)
    
    def _extract_phones(self, text: str) -> List[str]:
        """Extraction téléphones"""
        phone_patterns = [
            r'\b(?:\+33|0)[1-9](?:[.\-\s]?\d{2}){4}\b',
            r'\b\d{2}[\.\-\s]?\d{2}[\.\-\s]?\d{2}[\.\-\s]?\d{2}[\.\-\s]?\d{2}\b'
        ]
        
        phones = []
        for pattern in phone_patterns:
            phones.extend(re.findall(pattern, text))
        
        return list(set(phones))
    
    def _detect_sector(self, skills: List[str]) -> Optional[str]:
        """Détection secteur principal"""
        sector_scores = {}
        
        for sector, skills_list in self.skills_database.items():
            score = sum(1 for skill in skills if skill in skills_list)
            if score > 0:
                sector_scores[sector] = score
        
        if sector_scores:
            return max(sector_scores, key=sector_scores.get).title()
        
        return None
    
    def test_single_file(self, file_path: str, detailed: bool = True):
        """Test d'un seul fichier avec analyse détaillée"""
        print(f"\n🔍 TEST FICHIER: {file_path}")
        print("-" * 50)
        
        # 1. Extraction du texte
        extraction_result = self.test_text_extraction(file_path)
        
        print(f"📁 Fichier: {extraction_result['file']}")
        print(f"📏 Taille: {extraction_result['size']:,} octets")
        print(f"🔧 Méthode: {extraction_result['method']}")
        print(f"⏱️ Temps: {extraction_result['processing_time']:.1f}ms")
        
        if extraction_result['error']:
            print(f"❌ Erreur: {extraction_result['error']}")
            return
        
        text = extraction_result['text']
        print(f"📝 Texte extrait: {len(text):,} caractères")
        
        if len(text.strip()) == 0:
            print("⚠️ AUCUN TEXTE EXTRAIT - Possible PDF scanné nécessitant OCR")
            return
        
        # 2. Aperçu du texte
        if detailed:
            print(f"\n📖 APERÇU DU TEXTE (100 premiers caractères):")
            print(f"'{text[:100].strip()}...'")
        
        # 3. Analyse des données
        print(f"\n🧠 ANALYSE DES DONNÉES:")
        analysis = self.analyze_extracted_data(text)
        
        print(f"   📊 Statistiques texte:")
        print(f"      • Lignes: {analysis['lines_count']:,}")
        print(f"      • Mots: {analysis['words_count']:,}")
        
        print(f"   👤 Données candidat:")
        print(f"      • Nom: {analysis['name'] or 'Non détecté'}")
        print(f"      • Secteur: {analysis['sector'] or 'Non détecté'}")
        print(f"      • Expérience: {analysis['experience_years']} ans")
        
        print(f"   🎓 Compétences ({len(analysis['skills'])}):")
        for skill in analysis['skills'][:10]:  # Top 10
            print(f"      • {skill}")
        if len(analysis['skills']) > 10:
            print(f"      ... et {len(analysis['skills']) - 10} autres")
        
        if analysis['education']:
            print(f"   🎓 Formation: {', '.join(analysis['education'])}")
        
        if analysis['languages']:
            print(f"   🌍 Langues: {', '.join(analysis['languages'])}")
        
        if analysis['emails']:
            print(f"   📧 Emails: {', '.join(analysis['emails'])}")
        
        if analysis['phones']:
            print(f"   📞 Téléphones: {', '.join(analysis['phones'])}")
        
        # 4. Évaluation qualité
        print(f"\n📊 ÉVALUATION QUALITÉ:")
        quality_score = self._evaluate_quality(analysis)
        print(f"   Score global: {quality_score}/100")
        
        self.results.append({
            'file': file_path,
            'extraction': extraction_result,
            'analysis': analysis,
            'quality_score': quality_score
        })
    
    def _evaluate_quality(self, analysis: Dict) -> int:
        """Évaluation qualité de l'extraction"""
        score = 0
        
        # Texte extrait (30 points)
        if analysis['text_length'] > 100:
            score += 30
        elif analysis['text_length'] > 50:
            score += 15
        
        # Données candidat (40 points)
        if analysis['name']:
            score += 15
        if analysis['experience_years'] > 0:
            score += 10
        if analysis['skills']:
            score += 15
        
        # Données supplémentaires (30 points)
        if analysis['education']:
            score += 10
        if analysis['emails']:
            score += 10
        if analysis['languages']:
            score += 10
        
        return min(100, score)
    
    def test_directory(self, directory_path: str):
        """Test de tous les fichiers d'un dossier"""
        directory = Path(directory_path)
        
        if not directory.exists():
            print(f"❌ Dossier non trouvé: {directory_path}")
            return
        
        # Recherche des fichiers supportés
        supported_extensions = ['.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg']
        files = []
        
        for ext in supported_extensions:
            files.extend(directory.glob(f"*{ext}"))
            files.extend(directory.glob(f"*{ext.upper()}"))
        
        if not files:
            print(f"❌ Aucun fichier supporté trouvé dans: {directory_path}")
            return
        
        print(f"\n📁 TEST DOSSIER: {directory_path}")
        print(f"📊 Fichiers trouvés: {len(files)}")
        
        for i, file_path in enumerate(files, 1):
            print(f"\n{'='*20} FICHIER {i}/{len(files)} {'='*20}")
            self.test_single_file(str(file_path), detailed=False)
    
    def print_summary(self):
        """Résumé des tests"""
        print("\n" + "="*70)
        print("📊 RÉSUMÉ DES TESTS DE PARSING")
        print("="*70)
        
        print(f"📁 Fichiers testés: {self.stats['total_files']}")
        print(f"   • PDF: {self.stats['pdf_count']}")
        print(f"   • DOCX: {self.stats['docx_count']}")
        print(f"   • Images: {self.stats['image_count']}")
        print(f"   • TXT: {self.stats['txt_count']}")
        
        print(f"\n📊 Résultats:")
        print(f"   ✅ Succès: {self.stats['successful_extractions']}")
        print(f"   📄 Vides: {self.stats['empty_extractions']}")
        print(f"   ❌ Échecs: {self.stats['failed_extractions']}")
        
        if self.stats['total_files'] > 0:
            success_rate = (self.stats['successful_extractions'] / self.stats['total_files']) * 100
            print(f"\n🎯 Taux de succès: {success_rate:.1f}%")
        
        # Recommandations
        print(f"\n💡 RECOMMANDATIONS:")
        
        if self.stats['empty_extractions'] > 0:
            print("   • Fichiers vides détectés - Vérifiez si ce sont des PDF scannés")
            print("   • Installez tesseract pour l'OCR: brew install tesseract (macOS)")
        
        if not PDF_AVAILABLE:
            print("   • Installez PyPDF2: pip install PyPDF2")
        
        if not DOCX_AVAILABLE:
            print("   • Installez python-docx: pip install python-docx")
        
        if not OCR_AVAILABLE:
            print("   • Installez OCR: pip install pytesseract pillow")
        
        print("="*70)

def main():
    """Fonction principale"""
    tester = ParseTester()
    tester.print_header()
    
    if len(sys.argv) < 2:
        print("\n🚀 UTILISATION:")
        print("   python test_parsing.py <fichier>          # Test d'un fichier")
        print("   python test_parsing.py <dossier>          # Test d'un dossier")
        print("\n📝 EXEMPLES:")
        print("   python test_parsing.py cv_test.pdf")
        print("   python test_parsing.py ~/Desktop/CV\\ TEST/")
        print("   python test_parsing.py ~/Desktop/FDP\\ TEST/")
        return
    
    target_path = sys.argv[1]
    
    # Expansion du tilde (~)
    if target_path.startswith('~'):
        target_path = os.path.expanduser(target_path)
    
    if os.path.isfile(target_path):
        # Test d'un seul fichier
        tester.test_single_file(target_path)
    elif os.path.isdir(target_path):
        # Test d'un dossier
        tester.test_directory(target_path)
    else:
        print(f"❌ Chemin non valide: {target_path}")
        return
    
    tester.print_summary()

if __name__ == "__main__":
    main()
